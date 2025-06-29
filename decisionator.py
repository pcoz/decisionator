# decisionator.py
# Copyright (c) 2025 Edward Chalk (edward@fleetingswallow.com)
#
# Permission is hereby granted, free of charge, to any person obtaining a copy
# of this software and associated documentation files (the "Software"), to deal
# in the Software without restriction, including without limitation the rights
# to use, copy, modify, merge, publish, distribute, sublicense, and/or sell
# copies of the Software, and to permit persons to whom the Software is
# furnished to do so, subject to the following conditions:
#
# The above copyright notice and this permission notice shall be included in all
# copies or substantial portions of the Software.
#
# THE SOFTWARE IS PROVIDED "AS IS", WITHOUT WARRANTY OF ANY KIND, EXPRESS OR
# IMPLIED, INCLUDING BUT NOT LIMITED TO THE WARRANTIES OF MERCHANTABILITY,
# FITNESS FOR A PARTICULAR PURPOSE AND NONINFRINGEMENT. IN NO EVENT SHALL THE
# AUTHORS OR COPYRIGHT HOLDERS BE LIABLE FOR ANY CLAIM, DAMAGES OR OTHER
# LIABILITY, WHETHER IN AN ACTION OF CONTRACT, TORT OR OTHERWISE, ARISING FROM,
# OUT OF OR IN CONNECTION WITH THE SOFTWARE OR THE USE OR OTHER DEALINGS IN THE
# SOFTWARE.
#
# Attribution: This software may be used for commercial and non-commercial
# purposes, provided that credit is given to the original author:
# Edward Chalk (edward@fleetingswallow.com)
#
# Transactional Analysis Multiple Decision-Making Models Integrator
#
# Implementation with OpenAI API and workflow controller

#!/usr/bin/env python3
import os
import json
import openai
import sys
import re
import time
from typing import Dict, List, Any, Optional, Tuple, Callable, Union
from dataclasses import dataclass, asdict, field

from collections import defaultdict, Counter
from difflib import SequenceMatcher

from enum import Enum
import numpy as np

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

import markdown

import logging
import datetime
import itertools

from docx.shared import Inches

##Detect debug mode
import argparse

DEBUG_MODE = False
if __name__ == "__main__":
    parser = argparse.ArgumentParser()
    parser.add_argument("--debug", action="store_true", help="Enable debug logging")
    args, unknown = parser.parse_known_args()
    DEBUG_MODE = args.debug


####################################################################################################
# Helper functions for Markdown and Word document operations
####################################################################################################
def is_markdown(text):
    '''
    Detects if the supplied text includes Markdown formatting cues.
    Returns True if so, else False.
    '''
    return any(s in text for s in ('# ', '## ', '### ', '**', '`', '* '))



class DocAssembler:
    '''
    Utility class for building up a Word document section by section.
    Supports headings, paragraphs, tables, page breaks, and saving.
    Used throughout to build the final decision report.
    '''
    def __init__(self, title="AI Decision-Making Report"):
        self.doc = Document()
        self.section = "main"
        self.doc.add_heading(title, 0)
        self.doc.add_paragraph("Generated using TA Models and OpenAI.")
        self.doc.add_paragraph()
    
    def _add_markdown_to_doc(self, markdown_text):
        """
        Entry point for adding markdown text to the document. Splits the markdown
        into logical blocks and processes each.
        """
        text = self._normalize_text(markdown_text)
        blocks = self._split_into_blocks(text)
        for block in blocks:
            self._parse_block(block)

    def add_markdown_to_cell(self, cell, markdown_text):
        """
        Renders markdown-formatted text into a Word table cell.
        Supports bullets, bold, italics, etc., using the main markdown parsing logic.
        """
        # Use your main markdown splitting/parsing logic here:
        blocks = self._split_into_blocks(self._normalize_text(markdown_text))
        for block in blocks:
            if block['type'] == 'list':
                self._parse_list(block['content'], paragraph=None, container=cell)
            else:
                p = cell.add_paragraph()
                self._parse_inline_formatting(p, block['content'])
    

    def set_section(self, section_name):
        self.section = section_name
    
    def add_heading(self, text, level=1):
        self.doc.add_heading(text, level=level)
    
    def add_paragraph(self, text, style=None):
        self.doc.add_paragraph(text, style=style)
    
    def add_markdown(self, markdown_text):
        self._add_markdown_to_doc(markdown_text)
    
    def add_table(self, rows, cols):
        return self.doc.add_table(rows=rows, cols=cols)
    
    def add_page_break(self):
        self.doc.add_page_break()
    
    def save(self, filename):
        self.doc.save(filename)
    
    def get_doc(self):
        return self.doc
    
    def _normalize_text(self, text):
        """Normalize text by handling different line endings and excessive whitespace."""
        # Convert different line endings to \n
        text = re.sub(r'\r\n|\r', '\n', text)
        
        # Remove trailing whitespace from lines
        lines = [line.rstrip() for line in text.split('\n')]
        
        return '\n'.join(lines)
    
    def _split_into_blocks(self, text):
        """Split text into logical blocks (paragraphs, code blocks, lists, etc.)."""
        blocks = []
        lines = text.split('\n')
        current_block = []
        in_code_block = False
        code_block_lang = None
        in_table = False
        
        i = 0
        while i < len(lines):
            line = lines[i]
            stripped = line.strip()
            
            # Handle code blocks
            if stripped.startswith('```'):
                if in_code_block:
                    # End of code block
                    current_block.append(line)
                    blocks.append({
                        'type': 'code_block',
                        'content': '\n'.join(current_block),
                        'language': code_block_lang
                    })
                    current_block = []
                    in_code_block = False
                    code_block_lang = None
                else:
                    # Start of code block
                    if current_block:
                        blocks.append({'type': 'paragraph', 'content': '\n'.join(current_block)})
                        current_block = []
                    
                    code_block_lang = stripped[3:].strip() if len(stripped) > 3 else None
                    current_block.append(line)
                    in_code_block = True
                i += 1
                continue
            
            if in_code_block:
                current_block.append(line)
                i += 1
                continue
            
            # Handle tables
            if '|' in stripped and stripped.count('|') >= 2:
                if not in_table:
                    if current_block:
                        blocks.append({'type': 'paragraph', 'content': '\n'.join(current_block)})
                        current_block = []
                    in_table = True
                current_block.append(line)
            elif in_table and stripped == '':
                # End of table
                blocks.append({'type': 'table', 'content': '\n'.join(current_block)})
                current_block = []
                in_table = False
            elif in_table:
                # End of table due to non-table content
                blocks.append({'type': 'table', 'content': '\n'.join(current_block)})
                current_block = [line]
                in_table = False
            
            # Handle lists
            elif self._is_list_item(stripped):
                if current_block and not self._is_list_item(current_block[-1].strip()):
                    blocks.append({'type': 'paragraph', 'content': '\n'.join(current_block)})
                    current_block = []
                current_block.append(line)
            
            # Handle headers
            elif stripped.startswith('#'):
                if current_block:
                    blocks.append(self._determine_block_type('\n'.join(current_block)))
                    current_block = []
                blocks.append({'type': 'header', 'content': line})
            
            # Handle horizontal rules
            elif re.match(r'^[-*_]{3,}$', stripped):
                if current_block:
                    blocks.append(self._determine_block_type('\n'.join(current_block)))
                    current_block = []
                blocks.append({'type': 'hr', 'content': line})
            
            # Handle blockquotes
            elif stripped.startswith('>'):
                if current_block and not current_block[-1].strip().startswith('>'):
                    blocks.append(self._determine_block_type('\n'.join(current_block)))
                    current_block = []
                current_block.append(line)
            
            # Handle empty lines
            elif stripped == '':
                if current_block:
                    blocks.append(self._determine_block_type('\n'.join(current_block)))
                    current_block = []
            
            # Regular content
            else:
                current_block.append(line)
            
            i += 1
        
        # Handle remaining content
        if current_block:
            if in_table:
                blocks.append({'type': 'table', 'content': '\n'.join(current_block)})
            else:
                blocks.append(self._determine_block_type('\n'.join(current_block)))
        
        return blocks
    
    def _determine_block_type(self, content):
        """Determine the type of a content block."""
        stripped = content.strip()
        
        if self._is_list_block(content):
            return {'type': 'list', 'content': content}
        elif stripped.startswith('>'):
            return {'type': 'blockquote', 'content': content}
        else:
            return {'type': 'paragraph', 'content': content}
    
    def _is_list_item(self, line):
        """Check if a line is a list item."""
        # Unordered list patterns (including non-standard bullets)
        if re.match(r'^\s*[-*+·•‣⁃]\s+', line):
            return True
        # Ordered list patterns
        if re.match(r'^\s*\d+\.\s+', line):
            return True
        return False
    
    def _is_list_block(self, content):
        """Check if content block contains list items."""
        lines = content.split('\n')
        list_lines = sum(1 for line in lines if self._is_list_item(line.strip()))
        return list_lines > 0
    
    def _get_list_indent_level(self, line):
        """Calculate the indentation level of a list item."""
        # Count leading spaces/tabs before the list marker
        leading_spaces = len(line) - len(line.lstrip())
        # Convert tabs to spaces (assuming 4 spaces per tab)
        line_with_spaces = line.expandtabs(4)
        leading_spaces_normalized = len(line_with_spaces) - len(line_with_spaces.lstrip())
        
        # Each indentation level is typically 2 or 4 spaces
        # We'll use 2 spaces as the base unit
        return leading_spaces_normalized // 2
    
    def _parse_block(self, block):
        """Parse a single block based on its type."""
        block_type = block.get('type', 'paragraph')
        content = block.get('content', '')
        
        try:
            if block_type == 'header':
                self._parse_header(content)
            elif block_type == 'code_block':
                self._parse_code_block(content, block.get('language'))
            elif block_type == 'list':
                self._parse_list(content)
            elif block_type == 'table':
                self._parse_table(content)
            elif block_type == 'blockquote':
                self._parse_blockquote(content)
            elif block_type == 'hr':
                self._add_horizontal_rule()
            else:
                self._parse_paragraph(content)
        except Exception as e:
            logging.warning(f"Error parsing block type {block_type}: {str(e)}")
            # Fallback to plain paragraph
            self.doc.add_paragraph(content)
    
    def _parse_header(self, line):
        """Parse header line and add to document."""
        stripped = line.strip()
        if not stripped.startswith('#'):
            return
        
        # Count hash symbols
        level = 0
        for char in stripped:
            if char == '#':
                level += 1
            else:
                break
        
        # Limit to valid heading levels (1-6)
        level = min(max(level, 1), 6)
        
        # Extract header text
        header_text = stripped[level:].strip()
        
        if header_text:
            self.doc.add_heading(header_text, level=level)
    
    def _parse_code_block(self, content, language=None):
        """Parse code block and add to document."""
        lines = content.split('\n')
        
        # Remove opening and closing ```
        if lines and lines[0].strip().startswith('```'):
            lines = lines[1:]
        if lines and lines[-1].strip() == '```':
            lines = lines[:-1]
        
        # Add code paragraph
        for line in lines:
            p = self.doc.add_paragraph()
            run = p.add_run(line)
            run.font.name = 'Courier New'
            run.font.size = Pt(10)
            
            # Add light background for code blocks
            try:
                run.font.color.rgb = RGBColor(51, 51, 51)  # Dark gray text
            except:
                pass  # Ignore if color setting fails
    
    def _parse_list(self, content):
        """Parse list content and add to document with proper nesting support."""
        lines = content.split('\n')
        
        for line in lines:
            stripped = line.strip()
            if not stripped or not self._is_list_item(stripped):
                continue
            
            # Calculate indentation level
            indent_level = self._get_list_indent_level(line)
            
            # Extract list item text and determine list type
            if re.match(r'^\s*[-*+·•‣⁃]\s+', line):
                # Unordered list
                text = re.sub(r'^\s*[-*+·•‣⁃]\s+', '', line).strip()
                list_type = 'bullet'
            elif re.match(r'^\s*\d+\.\s+', line):
                # Ordered list
                text = re.sub(r'^\s*\d+\.\s+', '', line).strip()
                list_type = 'number'
            else:
                continue
            
            # Create paragraph with appropriate list style and indentation
            if list_type == 'bullet':
                if indent_level == 0:
                    p = self.doc.add_paragraph(style='List Bullet')
                elif indent_level == 1:
                    p = self.doc.add_paragraph(style='List Bullet 2')
                elif indent_level == 2:
                    p = self.doc.add_paragraph(style='List Bullet 3')
                else:
                    # For deeper nesting, use List Bullet 3 with additional left indent
                    p = self.doc.add_paragraph(style='List Bullet 3')
                    # Add extra indentation for very deep nesting
                    p.paragraph_format.left_indent = Inches(0.5 * (indent_level - 2))
            else:  # numbered list
                if indent_level == 0:
                    p = self.doc.add_paragraph(style='List Number')
                elif indent_level == 1:
                    p = self.doc.add_paragraph(style='List Number 2')
                elif indent_level == 2:
                    p = self.doc.add_paragraph(style='List Number 3')
                else:
                    # For deeper nesting, use List Number 3 with additional left indent
                    p = self.doc.add_paragraph(style='List Number 3')
                    # Add extra indentation for very deep nesting
                    p.paragraph_format.left_indent = Inches(0.5 * (indent_level - 2))
            
            # Clear any existing content in the paragraph
            p.clear()
            
            # Parse inline formatting including bold text at the beginning
            self._parse_inline_formatting(p, text)
    
    def _parse_table(self, content):
        """Parse table content and add to document."""
        lines = [line for line in content.split('\n') if line.strip()]
        if not lines:
            return
        
        # Filter out separator lines (lines with only |, -, and spaces)
        table_lines = []
        for line in lines:
            if not re.match(r'^\s*\|[\s\-\|]*\|\s*$', line):
                table_lines.append(line)
        
        if not table_lines:
            return
        
        # Parse first row to determine column count
        first_row = table_lines[0]
        cells = [cell.strip() for cell in first_row.split('|')]
        cells = [cell for cell in cells if cell]  # Remove empty cells from split
        
        if len(cells) < 2:
            # Not a valid table, treat as paragraph
            self._parse_paragraph(content)
            return
        
        # Create table
        table = self.doc.add_table(rows=len(table_lines), cols=len(cells))
        table.style = 'Table Grid'
        
        # Fill table
        for row_idx, line in enumerate(table_lines):
            cells_data = [cell.strip() for cell in line.split('|')]
            cells_data = [cell for cell in cells_data if cell]  # Remove empty
            
            for col_idx, cell_data in enumerate(cells_data):
                if col_idx < len(table.rows[row_idx].cells):
                    cell = table.rows[row_idx].cells[col_idx]
                    # Parse inline formatting in cell
                    self._parse_inline_formatting(cell.paragraphs[0], cell_data)
    
    def _parse_blockquote(self, content):
        """Parse blockquote content and add to document."""
        lines = content.split('\n')
        quote_text = []
        
        for line in lines:
            stripped = line.strip()
            if stripped.startswith('>'):
                quote_text.append(stripped[1:].strip())
            elif stripped == '':
                quote_text.append('')
            else:
                quote_text.append(stripped)
        
        # Add blockquote paragraph
        p = self.doc.add_paragraph()
        p.style = 'Quote'
        self._parse_inline_formatting(p, '\n'.join(quote_text))
    
    def _parse_paragraph(self, content):
        """Parse regular paragraph content."""
        if not content.strip():
            return
        
        lines = content.split('\n')
        for line in lines:
            if line.strip():
                p = self.doc.add_paragraph()
                self._parse_inline_formatting(p, line)
    
    def _parse_inline_formatting(self, paragraph, text):
        """Parse inline formatting like bold, italic, code, links."""
        if not text.strip():
            return
        
        # Handle various inline patterns with improved bold detection
        patterns = [
            (r'\*\*\*(.*?)\*\*\*', 'bold_italic'),  # Bold italic
            (r'\*\*(.*?)\*\*', 'bold'),             # Bold
            (r'\*(.*?)\*', 'italic'),               # Italic
            (r'~~(.*?)~~', 'strikethrough'),        # Strikethrough
            (r'`([^`]+)`', 'code'),                 # Inline code
            (r'\[([^\]]+)\]\(([^)]+)\)', 'link'),   # Links
        ]
        
        # Split text by patterns while preserving the matches
        parts = [text]
        format_info = [None]
        
        for pattern, format_type in patterns:
            new_parts = []
            new_format_info = []
            
            for i, part in enumerate(parts):
                if format_info[i] is not None:
                    # Already formatted, don't process further
                    new_parts.append(part)
                    new_format_info.append(format_info[i])
                    continue
                
                matches = list(re.finditer(pattern, part))
                if not matches:
                    new_parts.append(part)
                    new_format_info.append(None)
                    continue
                
                last_end = 0
                for match in matches:
                    # Add text before match
                    if match.start() > last_end:
                        new_parts.append(part[last_end:match.start()])
                        new_format_info.append(None)
                    
                    # Add formatted text
                    if format_type == 'link':
                        new_parts.append((match.group(1), match.group(2)))
                        new_format_info.append('link')
                    else:
                        new_parts.append(match.group(1))
                        new_format_info.append(format_type)
                    
                    last_end = match.end()
                
                # Add remaining text
                if last_end < len(part):
                    new_parts.append(part[last_end:])
                    new_format_info.append(None)
            
            parts = new_parts
            format_info = new_format_info
        
        # Add formatted runs to paragraph
        for part, fmt in zip(parts, format_info):
            if not part:
                continue
            
            if fmt == 'link':
                # Handle links
                text_part, url = part
                run = paragraph.add_run(text_part)
                try:
                    # Add hyperlink (requires more complex XML manipulation)
                    self._add_hyperlink(paragraph, run, url, text_part)
                except:
                    # Fallback: just add the text
                    run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
            else:
                run = paragraph.add_run(str(part))
                
                # Apply formatting
                if fmt == 'bold':
                    run.bold = True
                elif fmt == 'italic':
                    run.italic = True
                elif fmt == 'bold_italic':
                    run.bold = True
                    run.italic = True
                elif fmt == 'strikethrough':
                    run.font.strike = True
                elif fmt == 'code':
                    run.font.name = 'Courier New'
                    run.font.size = Pt(10)
                    try:
                        run.font.color.rgb = RGBColor(199, 37, 78)  # Red color for code
                    except:
                        pass
    
    def _add_hyperlink(self, paragraph, run, url, text):
        """Add hyperlink to paragraph (simplified version)."""
        # This is a simplified version - full hyperlink support requires more complex XML
        run.font.color.rgb = RGBColor(0, 0, 255)  # Blue color
        run.underline = True
    
    def _add_horizontal_rule(self):
        """Add horizontal rule to document."""
        p = self.doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run = p.add_run('─' * 50)
        run.font.color.rgb = RGBColor(128, 128, 128)  # Gray color

def style_table(table):
    """
    Applies black borders to all cells, shades the header row, bolds header text,
    and sets column widths: wide for 'Consideration', narrow for 'Type'/'Score'.
    """
    
    # 1. Set table-level border style first
    table.style = 'Table Grid'
    
    # 2. Apply black borders to all cells using a more direct approach
    for row in table.rows:
        for cell in row.cells:
            # Set borders using the paragraph format approach
            tc = cell._tc
            tcPr = tc.get_or_add_tcPr()
            
            # Remove any existing borders
            for borders in tcPr.xpath('.//w:tcBorders'):
                tcPr.remove(borders)
            
            # Add comprehensive border styling
            borders_xml = '''
            <w:tcBorders xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">
                <w:top w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:left w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:bottom w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:right w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideH w:val="single" w:sz="4" w:space="0" w:color="000000"/>
                <w:insideV w:val="single" w:sz="4" w:space="0" w:color="000000"/>
            </w:tcBorders>'''
            
            tcPr.append(parse_xml(borders_xml.strip()))

    # 2. Shade and bold header row
    hdr_cells = table.rows[0].cells
    for cell in hdr_cells:
        # Bold all text in header cells
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.bold = True
            # If no runs exist, create one and make it bold
            if not paragraph.runs:
                run = paragraph.runs[0] if paragraph.runs else paragraph.add_run(paragraph.text)
                run.font.bold = True
        
        # Add shading to header cells
        cell._tc.get_or_add_tcPr().append(
            parse_xml(f'<w:shd {nsdecls("w")} w:fill="D9D9D9"/>')
        )

    # 3. Set column widths (tight for Type/Score)
    if len(hdr_cells) == 3:
        # Set widths for header row
        hdr_cells[0].width = Inches(5.1)  # Consideration (adjust as needed)
        hdr_cells[1].width = Inches(0.65)  # Type (as narrow as practical)
        hdr_cells[2].width = Inches(0.65)  # Score (as narrow as practical)
        
        # Set widths for all other rows
        for row in table.rows:
            if len(row.cells) >= 3:
                row.cells[0].width = Inches(5.1)
                row.cells[1].width = Inches(0.65)
                row.cells[2].width = Inches(0.65)

from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml

def get_heatmap_color(score, min_neg, max_pos, type_):
    '''
    Maps a numerical score and consideration type to a color for table cell shading.
    
    Color logic based on SCORE SIGN, not consideration type:
    - Any consideration with negative score: Red gradient (light to dark red)
    - Any consideration with positive score: Green gradient (light to dark green)
    - Zero scores: No coloring
    
    This ensures intuitive coloring where negative scores are always red
    and positive scores are always green, regardless of the consideration type.
    '''
    try:
        s = float(score)
    except Exception:
        return None
    
    # Color based purely on score sign, not consideration type
    if s < 0:
        # Any negative score gets red coloring
        if min_neg < 0:
            t = abs(s) / abs(min_neg) if min_neg != 0 else 0
            r = int(251 + (255-251)*t)  # 251 -> 255 (light red to dark red)
            g = int(234 - 234*t)        # 234 -> 0 (reduce green)
            b = int(234 - 234*t)        # 234 -> 0 (reduce blue)
            return f"{r:02x}{g:02x}{b:02x}"
    
    elif s > 0:
        # Any positive score gets green coloring
        if max_pos > 0:
            t = s / max_pos if max_pos != 0 else 0
            r = int(245 - 245*t)        # 245 -> 0 (reduce red)
            g = int(255 - (255-204)*t)  # 255 -> 204 (light green to darker green)
            b = int(245 - 245*t)        # 245 -> 0 (reduce blue)
            return f"{r:02x}{g:02x}{b:02x}"
    
    # Zero scores or invalid ranges return no color
    return None

def normalized_score(cons):
    '''
    Returns a normalized (possibly negated) score for a consideration,
    so negative-oriented types always appear as negative numbers.
    '''
    try:
        val = float(cons.score)
    except Exception:
        val = 0.0
    if cons.type == "negative" and val > 0:
        return -val
    return val

# Sorting utility: sorts by score, then text alphabetically (case-insensitive)
def sort_considerations(conslist, negative=True):
    '''
    Sorts a list of Consideration objects.
    For negatives, sorts from most to least negative; for positives, ascending order.
    '''
    return sorted(
        conslist,
        key=lambda x: (float(x.score), str(x.text).lower())
    ) if negative else sorted(
        conslist,
        key=lambda x: (float(x.score), str(x.text).lower())
    )

# Deduping utility for deduping very similar "considerations" in the considerations tables.
# Increase the value of similarity_cutoff (in increments of 0.01) to increase the number of deduped considerations.
from difflib import SequenceMatcher


def signed_score(type_, score):
    '''
    Ensures negative considerations have negative scores and positive/avoidance ones have positive.
    Used to enforce type-appropriate sign in all tables.
    '''
    try:
        val = float(score)
    except Exception:
        val = 0.0
    # Treat 'avoidance' as 'positive'
    orientation = type_
    if type_ == "avoidance":
        orientation = "positive"
    if orientation == "negative" and val > 0:
        return -val
    if orientation == "positive" and val < 0:
        return abs(val)
    return val

####################################################################################################
# Helper Functions: Group & Amalgamate
####################################################################################################
def dedupe_considerations(conslist, similarity_cutoff=0.95):
    '''
    Deduplicate a list of Consideration objects, both by exact text and near-match (fuzzy).
    - similarity_cutoff: float between 0 and 1.0; higher = stricter.
    Returns a list with duplicates removed.
    '''
    seen = set()
    deduped = []
    texts = []
    for c in conslist:
        text = str(c.text).strip()
        # Exact dedupe
        if text.lower() in seen:
            continue
        # Fuzzy dedupe against previous
        is_dupe = False
        for prev in texts:
            if SequenceMatcher(None, text.lower(), prev.lower()).ratio() >= similarity_cutoff:
                is_dupe = True
                break
        if not is_dupe:
            deduped.append(c)
            texts.append(text)
            seen.add(text.lower())
    return deduped

##**Code for merging considerations**
class ConsiderationType(Enum):
    POSITIVE = "positive"
    NEGATIVE = "negative" 
    AVOIDANCE = "avoidance"

@dataclass
class MergedConsideration:
    """Data class for merged consideration results"""
    merged_text: str
    type: str
    score: float
    originals: List[Any]
    confidence: float = 1.0
    merge_method: str = "single"

##**
import uuid
import logging
from collections import defaultdict, Counter
from typing import List, Any, Dict, Optional

def merge_considerations_conceptually(
    conslist: List[Any], 
    controller: Any, 
    context_prompt: str = "",
    similarity_threshold: float = 0.8,  # Kept for legacy compatibility; ignored here
    max_group_size: int = 10,           # Kept for legacy compatibility; ignored here
    enable_semantic_grouping: bool = True,  # Kept for legacy compatibility; ignored here
    fallback_on_api_failure: bool = True,
    logger: Optional[logging.Logger] = None
) -> List[Dict[str, Any]]:
    """
    Iteratively merges considerations using the AI, always giving the AI the full current set.
    At each iteration, the AI can merge any (possibly overlapping) subset, returning:
      - Which IDs were merged,
      - The merged text.
    This continues until the AI says 'NO_MERGE'.
    All unmerged items are included as singles.

    Returns:
        List of merged consideration dicts, each with:
            merged_text, type, score, originals, confidence, merge_method, merged_ids
    """
    if not conslist:
        return []
    if logger is None:
        logger = logging.getLogger(__name__)

    print(f"\n🚀 STARTING CONSIDERATION CONSOLIDATION (Iterative AI, ID-driven)")
    logger.info(f"Processing {len(conslist)} considerations (AI iterative merging)")

    # ========== STEP 1: Validation & Bucketing ==========
    valid_types = {t.value for t in ConsiderationType}
    buckets = defaultdict(list)
    invalid_items = []

    # Attach a unique string ID to each consideration for AI tracking
    for i, consideration in enumerate(conslist):
        # Validate fields
        if not hasattr(consideration, 'text') or not hasattr(consideration, 'type'):
            logger.warning(f"Item {i} missing required attributes (text/type)")
            invalid_items.append(i)
            continue
        ctype = str(consideration.type).lower().strip()
        if ctype not in valid_types:
            logger.warning(f"Item {i} has invalid type '{ctype}', skipping")
            invalid_items.append(i)
            continue
        score = None
        try:
            score = float(getattr(consideration, 'score', 0.0))
        except Exception:
            score = 0.0
        text = str(consideration.text).strip()
        if not text:
            logger.warning(f"Item {i} has empty text, skipping")
            invalid_items.append(i)
            continue

        # Unique ID for tracking through the AI merge process
        if not hasattr(consideration, 'merge_id'):
            consideration.merge_id = str(uuid.uuid4())
        score_sign = "positive" if score >= 0 else "negative"
        bucket_key = (ctype, score_sign)
        # Attach index for error tracking
        consideration._original_index = i
        buckets[bucket_key].append(consideration)

    if invalid_items:
        print(f"⚠️ VALIDATION: Skipped {len(invalid_items)} invalid items (indices: {invalid_items})")
        logger.info(f"Skipped {len(invalid_items)} invalid items")

    if not buckets:
        print("❌ ERROR: No valid considerations found after validation")
        logger.warning("No valid considerations found after validation")
        return []

    print(f"📦 BUCKETING: Created {len(buckets)} type-based buckets")
    for (ctype, score_sign), bucket in buckets.items():
        print(f"   {ctype} ({score_sign}): {len(bucket)} items")
    logger.info(f"Created {len(buckets)} buckets for processing")

    # ========== STEP 2: AI-based Iterative Merging ==========

    merged_groups = []
    for (ctype, score_sign), bucket in buckets.items():
        print(f"\n🔍 PROCESSING BUCKET: {ctype} ({score_sign}) - {len(bucket)} items")
        logger.info(f"Processing bucket ({ctype}, {score_sign}) with {len(bucket)} items")

        # Maintain a working set for this bucket
        working = list(bucket)  # These are the "active" (unmerged) considerations for this bucket
        already_merged_ids = set()
        round_count = 0

        while len(working) > 1:
            round_count += 1
            # Build prompt for this merge round
            ai_prompt = _build_iterative_merge_prompt(
                working,
                context_prompt=context_prompt,
                bucket_type=ctype,
                bucket_score_sign=score_sign,
                logger=logger,
                round_number=round_count
            )
            logger.info(f"AI Merge Prompt (Round {round_count}):\n{ai_prompt}")

            try:
                ai_response = controller._call_openai_api(ai_prompt)
            except Exception as e:
                logger.error(f"AI merge call failed in round {round_count}: {e}")
                if fallback_on_api_failure:
                    break
                else:
                    raise

            # Parse AI response for merge instruction
            try:
                merge_instruction = _parse_ai_merge_response(ai_response, logger)
            except Exception as e:
                logger.error(f"Failed to parse AI response in round {round_count}: {e}")
                if fallback_on_api_failure:
                    break
                else:
                    raise

            # If NO_MERGE: exit loop, we're done with this bucket
            if merge_instruction.get("NO_MERGE", False):
                print(f"   ✅ AI indicated no further merges possible in round {round_count}.")
                break

            ids_to_merge = set(merge_instruction.get("ids", []))
            merged_text = merge_instruction.get("merged_text", "").strip()
            if not ids_to_merge or not merged_text:
                logger.warning(f"AI merge result incomplete or invalid (round {round_count}), skipping further merges.")
                break

            group_items = [item for item in working if getattr(item, "merge_id", None) in ids_to_merge]
            if not group_items:
                logger.warning(f"AI merge IDs did not match any current considerations in round {round_count}. Skipping merge.")
                break

            # Calculate average score, type majority, and originals for this merge
            avg_score = sum(float(getattr(item, "score", 0.0)) for item in group_items) / len(group_items)
            type_majority = Counter([str(getattr(item, "type", ctype)).lower().strip() for item in group_items]).most_common(1)[0][0]
            originals = [getattr(item, "original", item) for item in group_items]
            confidence = 0.95  # arbitrary, as AI merged

            # Add merged result to output
            merged_groups.append({
                "merged_text": merged_text,
                "type": type_majority,
                "score": avg_score,
                "originals": originals,
                "confidence": confidence,
                "merge_method": "ai_iterative",
                "merged_ids": list(ids_to_merge),
                "consolidation_count": len(group_items)
            })
            print(f"   🔄 CONSOLIDATED [{round_count}]: {len(group_items)} items → 1 by AI (IDs: {', '.join(ids_to_merge)})")
            logger.info(f"AI merged {len(group_items)} items (IDs: {ids_to_merge}) in round {round_count}")

            # Remove merged items from working set
            working = [item for item in working if getattr(item, "merge_id", None) not in ids_to_merge]
            already_merged_ids.update(ids_to_merge)

        # Add any remaining unmerged items as singles
        for item in working:
            merged_groups.append({
                "merged_text": str(getattr(item, "text", "")),
                "type": str(getattr(item, "type", ctype)),
                "score": float(getattr(item, "score", 0.0)),
                "originals": [getattr(item, "original", item)],
                "confidence": 1.0,
                "merge_method": "single",
                "merged_ids": [getattr(item, "merge_id", None)],
                "consolidation_count": 1
            })

    # ========== STEP 3: Final sorting, logging, reporting ==========

    # Sort by score ascending, regardless of type
    merged_groups = sorted(
        merged_groups,
        key=lambda g: float(g["score"])
    )

    # Logging and summary reporting as before
    total_originals = sum(len(group.get('originals', [])) for group in merged_groups)
    consolidations = [g for g in merged_groups if g.get('consolidation_count', 0) > 1]
    singles = [g for g in merged_groups if g.get('consolidation_count', 0) == 1]

    print(f"\n🎯 FINAL CONSOLIDATION RESULTS (Iterative AI):")
    print("=" * 50)
    print(f"📥 INPUT:  {len(conslist)} original considerations")
    print(f"📤 OUTPUT: {len(merged_groups)} final groups")
    print(f"🔄 CONSOLIDATED: {len(consolidations)} groups (representing {sum(g.get('consolidation_count', 0) for g in consolidations)} originals)")
    print(f"✅ SINGLES: {len(singles)} individual considerations")
    print(f"📊 REDUCTION: {((len(conslist) - len(merged_groups)) / len(conslist) * 100):.1f}% fewer items")
    if consolidations:
        print(f"\n🔍 CONSOLIDATION DETAILS:")
        for i, group in enumerate(consolidations):
            count = group.get('consolidation_count', 0)
            method = group.get('merge_method', 'unknown')
            ids = ', '.join(group.get('merged_ids', []))
            print(f"   Group {i+1}: {count} items (IDs: {ids}) → 1 (method: {method})")
    print("=" * 50)

    logger.info(f"Merge summary: {len(conslist)} input, {len(merged_groups)} output, {len(consolidations)} consolidated.")

    return merged_groups


def _build_iterative_merge_prompt(considerations, context_prompt, bucket_type, bucket_score_sign, logger, round_number):
    """
    Build a prompt for the AI to merge considerations, presenting all as an explicit ID+text list.
    """
    cons_str = ""
    for item in considerations:
        cons_str += f"- ID: {getattr(item, 'merge_id', '?')}\n"
        cons_str += f"  Type: {getattr(item, 'type', bucket_type)} | Score: {getattr(item, 'score', 0.0)}\n"
        cons_str += f"  Text: {getattr(item, 'text', '').strip()}\n"
    prompt = f'''
You are a technical decision analysis assistant.
Context: {context_prompt}

This is a list of considerations of type '{bucket_type}' ({bucket_score_sign}). Each has a unique 'ID'.

**Task:** For the list below, look for any two or more considerations that express the *same core idea or reasoning* (not just of the same type, but actually conceptually redundant or overlapping). If you find any such group, propose a single merged version as shown below.

Respond in **valid JSON** as:

{{
  "ids": ["ID1", "ID2", ...],    // the IDs you are merging (must be from the list below)
  "merged_text": "Merged text here, following best technical summary practices."
}}

**If, on this round, there are NO further mergeable groups, respond ONLY:**

```json
{{ "NO_MERGE": true }}
```
Here is the current list:
{cons_str}
'''
    logger.info(f"AI prompt for iterative merge (round {round_number}): {prompt[:1000]}...[truncated]" if len(prompt) > 1000 else prompt)
    return prompt


def _parse_ai_merge_response(response, logger):
    """
    Parse the AI's JSON output (which may be surrounded by Markdown code fences).
    """
    import json
    import re

    # Remove Markdown fences if present
    match = re.search(r'\{.*\}', response, re.DOTALL)
    if not match:
        # Try a less strict fallback
        logger.warning(f"Could not find JSON object in AI response. Raw response: {response}")
        if "NO_MERGE" in response:
            return {"NO_MERGE": True}
        raise ValueError("AI response did not contain a JSON object.")
    response_json = match.group(0)
    try:
        parsed = json.loads(response_json)
        return parsed
    except Exception as e:
        logger.error(f"JSON parse failed: {e}\nRaw: {response_json}")
        if "NO_MERGE" in response_json:
            return {"NO_MERGE": True}
        raise


def _sort_merged_groups(groups: List[Dict[str, Any]], logger: logging.Logger) -> List[Dict[str, Any]]:
    """Sort merged groups for optimal display order"""
    
    type_order = {"negative": 0, "avoidance": 1, "positive": 2}
    
    try:
        sorted_groups = sorted(
            groups, 
            key=lambda g: (
                type_order.get(g.get('type', 'unknown'), 99), 
                float(g.get('score', 0))
            )
        )
        return sorted_groups
        
    except Exception as e:
        logger.error(f"Failed to sort groups: {e}")
        return groups  # Return unsorted if sorting fails

##**
def _normalize_markdown_input(md: str, strip_html: bool, logger: logging.Logger) -> str:
    """Normalize and clean markdown input."""
    try:
        # Basic cleanup
        md_clean = md.strip()
        
        # Remove excessive whitespace
        md_clean = re.sub(r'\n\s*\n\s*\n+', '\n\n', md_clean)
        md_clean = re.sub(r'[ \t]+', ' ', md_clean)
        
        # Strip HTML tags if requested
        if strip_html:
            md_clean = re.sub(r'<[^>]+>', '', md_clean)
            # Clean up HTML entities
            html_entities = {
                '&amp;': '&', '&lt;': '<', '&gt;': '>', 
                '&quot;': '"', '&#39;': "'", '&nbsp;': ' '
            }
            for entity, char in html_entities.items():
                md_clean = md_clean.replace(entity, char)
        
        # Remove markdown formatting that might interfere
        #md_clean = re.sub(r'\*\*(.*?)\*\*', r'\1', md_clean)  # Bold
        #md_clean = re.sub(r'\*(.*?)\*', r'\1', md_clean)      # Italic
        #md_clean = re.sub(r'`(.*?)`', r'\1', md_clean)        # Inline code
        
        return md_clean
        
    except Exception as e:
        logger.error(f"Failed to normalize markdown input: {e}")
        return md.strip()  # Fallback to basic strip


def _remove_code_block_markers(lines: List[str], logger: logging.Logger) -> List[str]:
    """Remove code block markers from beginning and end."""
    try:
        if not lines:
            return lines
            
        # Remove opening code block
        if lines[0].strip().startswith('```'):
            lines = lines[1:]
            logger.debug("Removed opening code block marker")
        
        # Remove closing code block
        if lines and lines[-1].strip().endswith('```'):
            lines = lines[:-1]
            logger.debug("Removed closing code block marker")
            
        # Handle nested or multiple code blocks
        clean_lines = []
        in_code_block = False
        for line in lines:
            if line.strip().startswith('```'):
                in_code_block = not in_code_block
                continue
            if not in_code_block:
                clean_lines.append(line)
                
        return clean_lines
        
    except Exception as e:
        logger.warning(f"Error removing code block markers: {e}")
        return lines


def _find_bullet_start(lines: List[str]) -> Tuple[Optional[int], List[str]]:
    """
    Find the index of the first markdown bullet (of any recognized style),
    skipping any blank lines or code blocks before the list.
    - Handles code blocks, blank lines, and all standard and common Unicode bullets.
    - Returns (index of first bullet, bullet_patterns).
    """
    bullet_patterns = [
        r'^\s*[-•*+]\s+',           # Unordered (with optional indent)
        r'^\s*\d+\.\s+',            # Numbered (1., 2., etc)
        r'^\s*[a-zA-Z]\.\s+',       # Lettered (a., b., etc)
        r'^\s*→\s+',                # Unicode arrow
        r'^\s*▪\s+',                # Unicode square
        r'^\s*◦\s+',                # Unicode circle
        r'^\s*‣\s+',                # Unicode triangle bullet
        r'^\s*·\s+',                # Unicode middle dot
        r'^\s*⁃\s+',                # Unicode hyphen bullet
        r'^\s*‣\s+',                # Unicode triangle
        r'^\s*\(\d+\)\s+',          # (1), (2), ... styles
        r'^\s*\([a-zA-Z]\)\s+',     # (a), (b), ... styles
        r'^\s*—\s+',                # em dash
        r'^\s*•\s+',                # Bullet (duplicate for clarity)
    ]

    code_fence_open = False
    for i, line in enumerate(lines):
        # Remove trailing newlines/spaces and expand tabs for consistency
        line_expanded = line.expandtabs(4).rstrip('\n')
        stripped = line_expanded.strip()
        # Handle code blocks: never parse bullets inside
        if stripped.startswith("```"):
            code_fence_open = not code_fence_open
            continue
        if code_fence_open:
            continue
        # Skip blank lines
        if not stripped:
            continue
        # See if the line is a bullet
        for pattern in bullet_patterns:
            if re.match(pattern, stripped):
                return i, bullet_patterns
    return None, bullet_patterns


##Debug helper function
def debug_process_specific_md_line(md_line, logger=None):
    print("\n[DEBUG] Processing specific MD line:")
    print(f"  RAW: {repr(md_line)}")

    # Step 1: Normalization (as your markdown helpers do)
    norm = md_line.expandtabs(4).rstrip('\n').strip()
    print(f"  Normalized: {repr(norm)}")

    # Step 2: Bullet detection (patterns as per your _find_bullet_start)
    bullet_patterns = [
        r'^\s*[-•*+]\s+', r'^\s*\d+\.\s+', r'^\s*[a-zA-Z]\.\s+',
        r'^\s*→\s+', r'^\s*▪\s+', r'^\s*◦\s+', r'^\s*‣\s+',
        r'^\s*·\s+', r'^\s*⁃\s+', r'^\s*—\s+', r'^\s*•\s+',
        r'^\s*\(\d+\)\s+', r'^\s*\([a-zA-Z]\)\s+',
    ]
    matched = None
    for pattern in bullet_patterns:
        if re.match(pattern, norm):
            matched = pattern
            break
    print(f"  Bullet pattern matched: {matched}")

    # Step 3: Extract bullet content if matched
    if matched:
        match = re.match(matched, norm)
        bullet_content = norm[match.end():].strip()
        print(f"  Bullet content after marker: {repr(bullet_content)}")
    else:
        print("  No bullet pattern matched, treating as paragraph.")

    # Step 4: Inline formatting detection (bold/italic)
    # (Use your own parsing logic, here a quick scan)
    bold = re.findall(r'\*\*(.+?)\*\*', norm)
    italic = re.findall(r'\*(.+?)\*', norm)
    print(f"  Detected bold: {bold}")
    print(f"  Detected italic: {italic}")

    # Optional: Show how this would be added to a docx cell
    print("  Would be added as a bullet paragraph with parsed bold/italic in docx (if not in code block).")

    print("[DEBUG] Finished processing specific MD line.\n")
##

def _process_intro_section(lines: List[str], strip_html: bool, logger: logging.Logger) -> str:
    """Process introduction section before bullets."""
    try:
        if not lines:
            return ""
        
        # Join lines and clean up
        intro_parts = []
        for line in lines:
            line_clean = line.strip()
            if line_clean:
                intro_parts.append(line_clean)
        
        intro_text = ' '.join(intro_parts)
        
        # Additional cleanup for intro
        intro_text = re.sub(r'\s+', ' ', intro_text)  # Normalize whitespace
        intro_text = intro_text.strip()
        
        return intro_text
        
    except Exception as e:
        logger.error(f"Error processing intro section: {e}")
        return ""


def _extract_bullets(
    bullet_lines: List[str], 
    bullet_patterns: List[str], 
    max_length: int,
    strip_html: bool,
    logger: logging.Logger
) -> List[str]:
    """Extract and clean bullet points from lines."""
    bullets = []
    current_bullet = ""
    
    try:
        for line in bullet_lines:
            line_stripped = line.strip()
            if not line_stripped:
                continue
            
            # Check if this line starts a new bullet
            is_new_bullet = False
            bullet_text = ""
            
            for pattern in bullet_patterns:
                match = re.match(pattern, line_stripped)
                if match:
                    # Save previous bullet if exists
                    if current_bullet:
                        bullets.append(current_bullet.strip())
                    
                    # Extract bullet content
                    bullet_text = line_stripped[match.end():].strip()
                    current_bullet = bullet_text
                    is_new_bullet = True
                    break
            
            # If not a new bullet, might be continuation of previous
            if not is_new_bullet and current_bullet:
                # Handle multi-line bullets
                current_bullet += " " + line_stripped
            elif not is_new_bullet and line_stripped:
                # Standalone line that doesn't match bullet pattern
                logger.warning(f"Non-bullet line in bullet section: {line_stripped[:50]}...")
        
        # Don't forget the last bullet
        if current_bullet:
            bullets.append(current_bullet.strip())
        
        # Clean and validate bullets
        cleaned_bullets = []
        for bullet in bullets:
            bullet_clean = _clean_bullet_text(bullet, max_length, strip_html, logger)
            if bullet_clean:
                cleaned_bullets.append(bullet_clean)
        
        logger.debug(f"Extracted {len(cleaned_bullets)} bullets from {len(bullet_lines)} lines")
        return cleaned_bullets
        
    except Exception as e:
        logger.error(f"Error extracting bullets: {e}")
        return []


def _clean_bullet_text(bullet: str, max_length: int, strip_html: bool, logger: logging.Logger) -> str:
    """Clean individual bullet text."""
    try:
        if not bullet or not bullet.strip():
            return ""
        
        bullet_clean = bullet.strip()
        
        # Remove any remaining bullet markers that might have slipped through
        bullet_clean = re.sub(r'^[-•*+▪◦→]\s*', '', bullet_clean)
        bullet_clean = re.sub(r'^\d+\.\s*', '', bullet_clean)
        bullet_clean = re.sub(r'^[a-zA-Z]\.\s*', '', bullet_clean)
        
        # Length validation
        if len(bullet_clean) > max_length:
            logger.warning(f"Bullet text truncated from {len(bullet_clean)} to {max_length} chars")
            bullet_clean = bullet_clean[:max_length-3] + "..."
        
        # Final cleanup
        bullet_clean = re.sub(r'\s+', ' ', bullet_clean)
        bullet_clean = bullet_clean.strip()
        
        return bullet_clean
        
    except Exception as e:
        logger.error(f"Error cleaning bullet text: {e}")
        return bullet.strip() if bullet else ""


def _add_paragraph_safely(cell, text: str, style: str = None, logger: logging.Logger = None) -> bool:
    """Safely add paragraph to cell with error handling."""
    try:
        if not text or not text.strip():
            return False
        
        if style:
            cell.add_paragraph(text, style=style)
        else:
            cell.add_paragraph(text)
        
        return True
        
    except AttributeError as e:
        if logger:
            logger.error(f"Cell object missing add_paragraph method: {e}")
        raise RuntimeError(f"Invalid cell object: {e}")
    except Exception as e:
        if logger:
            logger.error(f"Failed to add paragraph to cell: {e}")
        raise RuntimeError(f"Failed to add content to cell: {e}")

####################################################################################################
# API helper functions
####################################################################################################
def get_option_text(opt: dict) -> str:
    """
    Safely returns the text label for a decision option, handling both 'text' and 'description' keys.
    Falls back to stringified dict if neither present.
    """
    return opt.get('text') or opt.get('description') or str(opt)

####################################################################################################
# Logging and API Config
####################################################################################################
# Set up logging to file for API responses and debug info.
LOG_FILENAME = "ta_api_calls.log"
logging.basicConfig(
    filename=LOG_FILENAME,
    filemode='a',
    format='%(asctime)s %(levelname)s: %(message)s',
    level=logging.INFO
)


def log_api_data(label: str, data: Any):
    '''
    Logs data to a log file with a label and ensures it is stringified.
    '''
    try:
        if isinstance(data, (dict, list)):
            text = json.dumps(data, indent=2, ensure_ascii=False)
        else:
            text = str(data)
        logging.info(f"[{label}] {text}")
    except Exception as e:
        logging.error(
            f"[{label}] (Logging Error) {e} // Original data: {data}")


# Read OpenAI API key from environment and configure API client.
OPENAI_API_KEY = os.getenv("OPENAI_API_KEY")
if not OPENAI_API_KEY:
    raise RuntimeError("Please set the OPENAI_API_KEY environment variable.")
openai.api_key = OPENAI_API_KEY

####################################################################################################
# Data field mapping and utility functions
####################################################################################################

# Dictionaries for mapping LLM numeric output fields to semantic labels.

EGO_FIELDS_BY_ID = {
    "ego_state": "ego_state",  # Name
    "1": "concerns",
    "2": "hopes",
    "3": "fears",
    "4": "score",
    "5": "reasoning"
}
SUBEGO_FIELDS_BY_ID = {
    "sub_state": "sub_state",
    "core_function": "core_function",
    "1": "concerns",
    "2": "stance",
    "3": "non_negotiables",
    "4": "reasoning"
}
MATRIX_FIELDS_BY_ID = {
    "ego_state": "ego_state",
    "maslow_level": "maslow_level",
    "1": "score",
    "2": "reasoning"
}

# Maps for field_id
FIELD_ID_TO_LABEL = {
    "1": "concern",
    "2": "hope",
    "3": "fear",
    "4": "score",  # not usually needed as consideration
    "5": "reasoning",  # will filter OUT for tables!
}
FIELD_ID_TO_ORIENTATION = {
    "1": "negative",   # Concerns
    "2": "positive",   # Hopes
    "3": "negative",   # Fears
    "4": "neutral",    # Scores, rarely needed as a consideration
    "5": "neutral",    # Reasoning – you’ll filter these out for tables!
}


@dataclass
class EgoStateResponse:
    '''
    Data container for storing results of an ego state analysis in Model 1.
    '''
    ego_state: str
    concerns: List[str]
    hopes: List[str]
    fears: List[str]
    score: float
    reasoning: str


@dataclass
class SubEgoStateResponse:
    '''
    Data container for results from a sub-ego state analysis in Model 2.
    '''
    sub_state: str
    core_function: str
    concerns: List[str]
    stance: int  # -2 to +2
    non_negotiables: List[str]
    reasoning: str


@dataclass
class MatrixCell:
    '''
    Stores a single (ego_state, maslow_level) cell from Model 3’s matrix.
    '''
    ego_state: str
    maslow_level: str
    score: float
    reasoning: str


@dataclass
class DecisionResult:
    '''
    Stores the main output/result from any of the three decision models.
    '''
    model_used: str
    recommendation: str
    confidence_score: float
    conditions: List[str]
    reasoning: str
    detailed_analysis: Dict[str, Any]
    
@dataclass
class Consideration:
    '''
    Represents a single consideration (pro/con/requirement/score)
    as output by a decision model, with metadata for sorting/grouping.
    '''
    text: str
    source_model: str
    source_context: str
    type: str  # 'positive', 'negative', or 'avoidance'
    option: str = None
    score: float = None
    extra: dict = field(default_factory=dict)
    field_id: str = None

def extract_numeric(val):
    '''
    Extracts a float from a string (or int/float), using only 0-9 and '.' characters.
    If conversion fails, returns 0.0.
    '''
    if isinstance(val, (int, float)):
        return float(val)
    if not isinstance(val, str):
        return 0.0
    digits = ''.join([c for c in val if (c.isdigit() or c == '.')])
    try:
        return float(digits) if digits else 0.0
    except Exception:
        return 0.0


def map_fields_by_id(data: dict, id_map: dict, dataclass_type):
    '''Convert LLM output (with numeric keys) to dataclass fields and handle type coercion.'''
    mapped = {}
    for id_key, attr in id_map.items():
        if id_key in data:
            mapped[attr] = data[id_key]
    # Type coercion for each field
    for f in dataclass_type.__dataclass_fields__:
        typ = dataclass_type.__dataclass_fields__[f].type
        if f in mapped:
            if typ == float:
                mapped[f] = extract_numeric(mapped[f])
            elif typ == int:
                mapped[f] = int(extract_numeric(mapped[f]))
            elif typ == list or typ == List[str]:
                v = mapped[f]
                if isinstance(v, str):
                    try:
                        mapped[f] = json.loads(v)
                    except Exception:
                        mapped[f] = [s.strip() for s in v.split(",")]
                elif not isinstance(v, list):
                    mapped[f] = [v]
    # If a required field is missing, fill with defaults
    for f in dataclass_type.__dataclass_fields__:
        typ = dataclass_type.__dataclass_fields__[f].type
        if f not in mapped:
            if typ == float:
                mapped[f] = 0.0
            elif typ == int: mapped[f] = 0
            elif typ == list or typ == List[str]:
                mapped[f] = []
            else: mapped[f] = ""
    return dataclass_type(**mapped)

####################################################################################################
# Core Enums and Controller Classes
####################################################################################################

class DecisionModel(Enum):
    '''
    Enum of supported Transactional Analysis decision models.
    Used to select which workflow to run.
    '''
    DEMOCRATIC_COUNCIL = "model1"
    SECOND_ORDER_NEGOTIATIONS = "model2"
    MASLOW_TA_MATRIX = "model3"


class WorkflowController:
    '''
    Orchestrates all operations for the three supported decision models.
    Handles prompt creation, OpenAI API calls, model execution, and parsing.
    Provides public methods for running each model and extracting options.
    '''

    def __init__(self):
        self.workflows = {
            DecisionModel.DEMOCRATIC_COUNCIL: {
                "steps": [
                    "independent_ego_analysis",
                    "council_synthesis",
                    "final_decision"
                ],
                "ego_states": ["Parent", "Adult", "Child"]
            },
            DecisionModel.SECOND_ORDER_NEGOTIATIONS: {
                "steps": [
                    "sub_ego_analysis",
                    "cluster_dialogues",
                    "cross_cluster_negotiation",
                    "weighted_vote"
                ],
                "sub_ego_states": [
                    "Parent-in-Parent", "Adult-in-Parent", "Child-in-Parent",
                    "Parent-in-Adult", "Adult-in-Adult", "Child-in-Adult",
                    "Parent-in-Child", "Adult-in-Child", "Child-in-Child"
                ]
            },
            DecisionModel.MASLOW_TA_MATRIX: {
                "steps": [
                    "matrix_evaluation",
                    "tier_by_tier_check",
                    "mitigation_planning",
                    "utility_calculation"
                ],
                "ego_states": ["Parent", "Adult", "Child"],
                "maslow_levels": ["Physiological", "Safety", "Love/Belonging", "Esteem", "Self-Actualization"]
            }
        }

    def extract_options(self, problem: str) -> list:
        prompt = f'''
    Given the following decision problem, enumerate the main mutually-exclusive options under consideration, in clear, concise terms.
    Return as JSON with unique short IDs and one-line descriptions.
    
    Problem:
    {problem}
    
    Example:
    {{
      "options": [
        {{"id": "A", "text": "Fish in the safe zone"}},
        {{"id": "B", "text": "Fish in the disputed zone"}},
        {{"id": "C", "text": "Repair the boat before going out"}}
      ]
    }}
    '''
        response = self._call_openai_api(prompt)
        data = self._handle_json_parse(response)
        return data.get("options", [])
    
    def _make_final_decision_model1(self, problem: str, ego_responses: List[EgoStateResponse], synthesis: Dict[str, Any]) -> DecisionResult:
        '''Make the final decision for Model 1'''
        avg_score = np.mean([response.score for response in ego_responses])
        # On -10 to +10 scale: Accept for positive, Reject for negative
        recommendation = "Accept" if avg_score >= 0 else "Reject"
        if synthesis["consensus_level"] < 3:
            recommendation += " (with reservations)"
        conditions = synthesis.get("mitigation_plans", [])
        return DecisionResult(
            model_used="Model 1: Democratic Ego State Council",
            recommendation=recommendation,
            confidence_score=float(avg_score),
            conditions=conditions,
            reasoning=synthesis["synthesis_reasoning"],
            detailed_analysis={
                "ego_responses": [asdict(r) for r in ego_responses],
                "synthesis": synthesis,
                "average_score": avg_score
            }
        )
    
    def _conduct_weighted_vote_model2(self, problem: str, negotiation_results: Dict[str, Any]) -> DecisionResult:
        '''
        Conduct the final weighted vote for Model 2 (with stances in -10 to +10 range).
        '''
        # Define weights for sub-ego states
        weights = {
            "Parent-in-Parent": 3, "Adult-in-Parent": 3, "Child-in-Parent": 3,
            "Parent-in-Adult": 2, "Adult-in-Adult": 2, "Child-in-Adult": 2,
            "Parent-in-Child": 1, "Adult-in-Child": 1, "Child-in-Child": 1
        }
        total_weight = sum(weights.values())
    
        # Try to extract the stances (should be a dict: sub_state -> stance)
        stances = {}
        # Try negotiation_results["final_stances"], fallback to ["stance"] in sub-states list
        if negotiation_results and "final_stances" in negotiation_results:
            stances = negotiation_results["final_stances"]
        elif negotiation_results:
            # Fallback: try to infer from negotiation_results (e.g., a list of sub-state dicts)
            for sub in weights:
                stance = negotiation_results.get(sub)
                if isinstance(stance, dict) and "stance" in stance:
                    stances[sub] = stance["stance"]
                elif isinstance(stance, (int, float)):
                    stances[sub] = stance
                # else: skip if not found
    
        # Fill in missing sub-ego stances with 0
        weighted_sum = 0
        for sub, weight in weights.items():
            stance = float(stances.get(sub, 0))  # Default to 0 if missing
            weighted_sum += stance * weight
    
        # Max/min possible sums
        max_possible = 10 * total_weight
        min_possible = -10 * total_weight
    
        # Normalized confidence: -1 (all -10) to +1 (all +10)
        if max_possible - min_possible != 0:
            confidence = (weighted_sum - min_possible) / (max_possible - min_possible) * 2 - 1
        else:
            confidence = 0
    
        recommendation = "Accept" if weighted_sum > 0 else "Reject"
    
        return DecisionResult(
            model_used="Model 2: Second-Order Ego State Negotiations",
            recommendation=recommendation,
            confidence_score=confidence,  # -1..+1 (or you can use abs(confidence) if you want only magnitude)
            conditions=negotiation_results.get(
                "agreed_actions", []) if negotiation_results else [],
            reasoning=f"Weighted sum of stances: {weighted_sum:.2f} (scale: {min_possible} to {max_possible}); Confidence: {confidence:.2f}",
            detailed_analysis={
                "negotiation_results": negotiation_results,
                "weighted_sum": weighted_sum,
                "max_possible": max_possible,
                "min_possible": min_possible,
                "confidence": confidence,
                "stances": stances,
                "weights": weights
            }
        )
    
    
    def _check_tiers_model3(self, problem: str, matrix_data: List[MatrixCell]) -> Tuple[Dict[str, Any], List[str]]:
        '''Check tiers bottom-up for Model 3'''
        levels = {}
        for cell in matrix_data:
            if cell.maslow_level not in levels:
                levels[cell.maslow_level] = []
            levels[cell.maslow_level].append(cell)
        tier_order = ["Physiological", "Safety",
            "Love/Belonging", "Esteem", "Self-Actualization"]
        tier_results = {}
        mitigations = []
        for tier in tier_order:
            if tier in levels:
                scores = [cell.score for cell in levels[tier]]
                min_score = min(scores)
                avg_score = np.mean(scores)
                tier_results[tier] = {
                    "scores": scores,
                    "min_score": min_score,
                    "avg_score": avg_score,
                    "passes": min_score >= 3.0
                }
                if min_score < 3.0:
                    mitigations.append(
                        f"Mitigation needed for {tier} (min score: {min_score})")
        return tier_results, mitigations
    
    def _calculate_utility_model3(self, problem: str, matrix_data: List[MatrixCell], tier_results: Dict[str, Any]) -> DecisionResult:
        '''Calculate final utility score for Model 3 (with scores in -10 to +10 range).'''
        tier_weights = {
            "Physiological": 1,
            "Safety": 2,
            "Love/Belonging": 3,
            "Esteem": 4,
            "Self-Actualization": 5
        }
        total_utility = 0
        max_possible = 0
        min_possible = 0
        for tier, weight in tier_weights.items():
            if tier in tier_results:
                avg_score = tier_results[tier]["avg_score"]
                total_utility += weight * avg_score
                max_possible += weight * 10
                min_possible += weight * -10
    
        # Normalized utility: 0 (all -10s), 0.5 (neutral/zero), 1 (all +10s)
        if max_possible - min_possible > 0:
            scaled_score = (total_utility - min_possible) / (max_possible - min_possible)
        else:
            scaled_score = 0.5  # Default to neutral if denominator is zero (shouldn't happen)
    
        # Accept if net utility is above neutral (scaled_score > 0.5)
        recommendation = "Accept" if scaled_score > 0.5 else "Reject"
    
        # Address any tiers with an average < 0 as "concern"
        conditions = []
        for tier, results in tier_results.items():
            if results["min_score"] < 0:
                conditions.append(
                    f"Address {tier} concerns (min score: {results['min_score']:.1f})"
                )
    
        return DecisionResult(
            model_used="Model 3: Maslow-TA Decision Matrix",
            recommendation=recommendation,
            confidence_score=scaled_score,  # 0 = all -10s, 0.5 = net neutral, 1 = all +10s
            conditions=conditions,
            reasoning=f"Weighted utility: {total_utility:.1f} (range: {min_possible} to {max_possible}); Scaled: {scaled_score:.2f}",
            detailed_analysis={
                "matrix_data": [asdict(cell) for cell in matrix_data],
                "tier_results": tier_results,
                "total_utility": total_utility,
                "max_possible": max_possible,
                "min_possible": min_possible,
                "scaled_score": scaled_score
            }
        )
    

    def _everyday_language_summary_prompt(self, problem: str, model_output: Dict[str, Any]) -> str:
        return f'''
You are an expert at translating technical or psychological reports into clear, plain, everyday language for normal people.
Below is a decision problem and the output of a decision-making model that uses psychology terms. Your job is to summarise, in simple language, what this output means, what the main recommendation is, and why – *without* using Transactional Analysis or psychological jargon.

Original Question:
{problem}

Model Output:
{json.dumps(model_output, indent=2)}

Please summarise in clear, everyday language for a general audience.
    '''
    

    def _is_error(self, data: dict):
        '''Check if a JSON object is an error response from the API.'''
        return "__error__" in data

    def _strip_json_code_block(self, text: str) -> str:
        """
        Extract just the JSON payload from an LLM reply that may contain
        Markdown fences or stray commentary.

        Strategy
        --------
        1. Look for a fenced block  ```json … ```  (or plain ``` … ```).
           If found, return the braces’ contents.
        2. Fallback: locate the first '{' and the last '}' in the reply.
           If that slice parses successfully with json.loads(), return it.
        3. If neither tactic succeeds, give the raw text back so the caller
           can surface a clear “JSON parse error” message.

        The function never raises; it only returns a string.
        """
        import re, json   # local import keeps the method standalone

        text = text.strip()

        # --- 1) fenced ```json``` or ``` block anywhere in the text ----------
        fence = re.search(r"```(?:json)?\s*({.*?})\s*```", text,
                          flags=re.S | re.I)
        if fence:
            return fence.group(1).strip()

        # --- 2) bare JSON object somewhere in the message --------------------
        first_brace = text.find('{')
        last_brace = text.rfind('}')
        first_bracket = text.find('[')
        last_bracket = text.rfind(']')
        # Prefer object if both present, else try array
        if 0 <= first_brace < last_brace:
            candidate = text[first_brace:last_brace + 1]
            try:
                json.loads(candidate)
                return candidate.strip()
            except Exception:
                pass
        if 0 <= first_bracket < last_bracket:
            candidate = text[first_bracket:last_bracket + 1]
            try:
                json.loads(candidate)
                return candidate.strip()
            except Exception:
                pass

        # --- 3) give up – let _handle_json_parse() raise a JSON error ---------
        return text

    def _handle_json_parse(self, response):
        '''
        Robustly parse JSON and halt on model-access errors.
        Always returns a dict. Also logs sanitized text and errors.
        '''
        sanitized = self._strip_json_code_block(response)
        log_api_data("API_SANITIZED", sanitized)
        try:
            data = json.loads(sanitized)
            log_api_data("API_PARSED", data)
            
            print("[DEBUG] Raw model response:", response)
            print("[DEBUG] Sanitized for JSON:", sanitized)

            return data
        except Exception as e:
            # Try to fix unescaped control characters
            def escape_control_chars(s):
                # This escapes literal line breaks and tabs inside strings
                def repl(m):
                    value = m.group(0)
                    return value.replace('\n', '\\n').replace('\r', '\\r').replace('\t', '\\t')
                # Only operate inside double-quoted strings
                return re.sub(r'(?<=")(.*?)(?=")', repl, s, flags=re.DOTALL)
            
            sanitized_fixed = escape_control_chars(sanitized)
            try:
                data = json.loads(sanitized_fixed)
                log_api_data("API_PARSED_FIXED", data)
                print("[WARN] Fixed unescaped control chars in JSON.")
                return data
            except Exception as e2:
                print("\nJSON parse error:", e2)
                print("Raw response was:\n", sanitized, "\n")
                log_api_data("API_JSON_ERROR", {
                             "exception": str(e2), "raw": sanitized})
                if "model" in sanitized and ("not available" in sanitized or "not found" in sanitized):
                    print(
                        "\n*** ERROR: The OpenAI API key you are using does not have access to the model you requested. ***")
                    print(
                        "-> Try a different model (such as gpt-3.5-turbo, gpt-4o, or check your account entitlements).")
                    print(
                        "-> See https://platform.openai.com/docs/models for current options, or upgrade your plan if needed.")
                    sys.exit(1)
                if "You do not have access" in sanitized:
                    print(
                        "\n*** ERROR: Your API key does not have access to this model. ***")
                    print("-> Try a different model or upgrade your account.")
                    sys.exit(1)
                return {"__error__": f"JSON parse error: {str(e2)}"}

    def _call_openai_api(self, prompt: str) -> str:
        '''Make API call to OpenAI v1.x API and handle model access errors. Also log every return.'''
        try:
            response = openai.chat.completions.create(
                # === MODEL OPTIONS (uncomment ONE) ===

                # GPT-4.1 Turbo (high quality, efficient)
                model="gpt-4-1106-preview",

                # model="gpt-4o",
                # model="gpt-3.5-turbo",

                messages=[
                    {"role": "system", "content": "You are an expert in Transactional Analysis and decision-making psychology. Always respond with valid JSON as requested, or with Markdown when asked for a summary."},
                    {"role": "user", "content": prompt}
                ],
                temperature=0.7,
                max_tokens=2000
            )
            content = response.choices[0].message.content.strip(
            ) if response.choices else ""
            # DEBUG: See what the model sends
            print(f"API raw response: {content!r}")
            log_api_data("API_RAW", content)
            return content
        except Exception as e:
            error_str = str(e)
            print(f"OpenAI API Error: {error_str}")
            log_api_data("API_ERROR", error_str)
            # Handle "model not available" or "not found" error explicitly
            if ("model" in error_str and "not available" in error_str) or \
               ("model" in error_str and "not found" in error_str) or \
               ("You do not have access" in error_str):
                print(
                    "\n*** ERROR: The OpenAI API key you are using does not have access to the model you requested. ***")
                print(
                    "-> Try a different model (such as gpt-3.5-turbo, gpt-4o, or check your account entitlements).")
                print(
                    "-> See https://platform.openai.com/docs/models for current options, or upgrade your plan if needed.")
                sys.exit(1)
            return json.dumps({"__error__": error_str})

    def _execute_model1(self, problem: str, options, filtered_pairs, doc, considerations, log_callback=None) -> DecisionResult:
        if log_callback:
            log_callback("start", {"model": "model1", "problem": problem})
        print("Executing Model 1: Democratic Ego State Council")
        ego_responses = []
        for ego_state in self.workflows[DecisionModel.DEMOCRATIC_COUNCIL]["ego_states"]:
            # Call API with the new prompt (returns a flat list of dicts)
            response = self._call_openai_api(self._ego_state_prompt(problem, ego_state, options, filtered_pairs))
            if log_callback:
                log_callback("ego_response", {"ego_state": ego_state, "raw_response": response})
            if doc:
                doc.add_heading(f"{ego_state} Ego State Analysis", level=3)
                try:
                    parsed = json.loads(self._strip_json_code_block(response))
                    if isinstance(parsed, list):
                        for item in parsed:
                            doc.add_paragraph(f"{item.get('type', '')}: {item.get('text', '')} (Option: {item.get('option', '')}, Score: {item.get('score', '')})")
                    else:
                        for k, v in parsed.items():
                            doc.add_paragraph(f"{k}: {v}")
                except Exception:
                    doc.add_paragraph(f"Raw: {response}")
    
            # Parse and extract considerations
            data = self._handle_json_parse(response)
            if self._is_error(data):
                print("API call failed for ego state:", ego_state, data["__error__"])
                ego_responses.append(
                    EgoStateResponse(
                        ego_state=ego_state,
                        concerns=[],
                        hopes=[],
                        fears=[],
                        score=0.0,
                        reasoning="API call failed: " + data["__error__"]
                    )
                )
            else:
                # data is expected to be a list of consideration dicts
                # (If not, skip gracefully)
                if isinstance(data, list):
                    for cons in data:
                        c = Consideration(
                            text=cons.get("text", ""),
                            source_model="model1",
                            source_context=ego_state,
                            type=cons.get("type", ""),
                            option=cons.get("option"),
                            score=cons.get("score"),
                        )
                        considerations.append(c)
                        if log_callback:
                            log_callback("consideration", asdict(c))
                # Optionally store an empty EgoStateResponse for backward compatibility
                ego_responses.append(
                    EgoStateResponse(
                        ego_state=ego_state,
                        concerns=[],  # deprecated in new format
                        hopes=[],
                        fears=[],
                        score=0.0,
                        reasoning=""
                    )
                )
    
        # Synthesis prompt (unchanged)
        synthesis_resp = self._call_openai_api(self._council_synthesis_prompt(problem, ego_responses))
        if log_callback:
            log_callback("council_synthesis", {"raw_response": synthesis_resp})
        if doc:
            doc.add_heading("Council Synthesis", level=3)
            try:
                parsed = json.loads(self._strip_json_code_block(synthesis_resp))
                for k, v in parsed.items():
                    doc.add_paragraph(f"{k}: {v}")
            except Exception:
                doc.add_paragraph(f"Raw: {synthesis_resp}")
        synthesis = self._handle_json_parse(synthesis_resp)
        if self._is_error(synthesis):
            print("API call failed for council synthesis:", synthesis["__error__"])
            synthesis = {
                "shared_themes": [],
                "tension_points": [],
                "mitigation_plans": [],
                "consensus_level": 0,
                "synthesis_reasoning": "API call failed: " + synthesis["__error__"]
            }
        if log_callback:
            log_callback("synthesis_parsed", synthesis)
        decision = self._make_final_decision_model1(problem, ego_responses, synthesis)
        if log_callback:
            log_callback("final_decision", asdict(decision))
    
        print(f"[DEBUG] considerations: {[asdict(c) for c in considerations]}")
    
        return decision

    def _execute_model2(self, problem: str, options, filtered_pairs, doc, considerations, log_callback=None) -> DecisionResult:
        if log_callback:
            log_callback("start", {"model": "model2", "problem": problem})
        print("Executing Model 2: Second-Order Ego State Negotiations")
        sub_ego_responses = []
        for sub_state in self.workflows[DecisionModel.SECOND_ORDER_NEGOTIATIONS]["sub_ego_states"]:
            response = self._call_openai_api(self._sub_ego_state_prompt(problem, sub_state, options, filtered_pairs))
            if log_callback:
                log_callback("sub_ego_response", {"sub_state": sub_state, "raw_response": response})
            if doc:
                doc.add_heading(f"{sub_state} Sub-Ego Analysis", level=3)
                try:
                    parsed = json.loads(self._strip_json_code_block(response))
                    if isinstance(parsed, list):
                        for item in parsed:
                            doc.add_paragraph(f"{item.get('type', '')}: {item.get('text', '')} (Option: {item.get('option', '')}, Score: {item.get('score', '')})")
                    else:
                        for k, v in parsed.items():
                            doc.add_paragraph(f"{k}: {v}")
                except Exception:
                    doc.add_paragraph(f"Raw: {response}")
    
            data = self._handle_json_parse(response)
            if self._is_error(data):
                print("API call failed for sub-ego state:", sub_state, data["__error__"])
                sub_ego_responses.append(
                    SubEgoStateResponse(
                        sub_state=sub_state,
                        core_function="API call failed",
                        concerns=[],
                        stance=0,
                        non_negotiables=[],
                        reasoning="API call failed: " + data["__error__"]
                    )
                )
            else:
                # data is expected to be a list of consideration dicts
                if isinstance(data, list):
                    for cons in data:
                        c = Consideration(
                            text=cons.get("text", ""),
                            source_model="model2",
                            source_context=sub_state,
                            type=cons.get("type", ""),
                            option=cons.get("option"),
                            score=cons.get("score"),
                        )
                        considerations.append(c)
                        if log_callback:
                            log_callback("consideration", asdict(c))
                # Optionally store an empty SubEgoStateResponse for legacy/future use
                sub_ego_responses.append(
                    SubEgoStateResponse(
                        sub_state=sub_state,
                        core_function=self._sub_ego_state_desc(sub_state),
                        concerns=[],  # deprecated
                        stance=0,
                        non_negotiables=[],
                        reasoning=""
                    )
                )
    
        # Handle cluster dialogues as before
        cluster_results_resp = self._call_openai_api(self._cluster_dialogues_prompt(problem, sub_ego_responses))
        if log_callback:
            log_callback("cluster_dialogues", {"raw_response": cluster_results_resp})
        if doc:
            doc.add_heading("Cluster Dialogues", level=3)
            try:
                parsed = json.loads(self._strip_json_code_block(cluster_results_resp))
                for k, v in parsed.items():
                    doc.add_paragraph(f"{k}: {v}")
            except Exception:
                doc.add_paragraph(f"Raw: {cluster_results_resp}")
        cluster_results = self._handle_json_parse(cluster_results_resp)
        if self._is_error(cluster_results):
            print("API call failed for cluster dialogues:", cluster_results["__error__"])
            cluster_results = {}
        if log_callback:
            log_callback("cluster_dialogues_parsed", cluster_results)
    
        # Cross-cluster negotiation
        negotiation_results_resp = self._call_openai_api(self._cross_cluster_negotiation_prompt(problem, cluster_results))
        if log_callback:
            log_callback("cross_cluster_negotiation", {"raw_response": negotiation_results_resp})
        if doc:
            doc.add_heading("Cross-Cluster Negotiation", level=3)
            try:
                parsed = json.loads(self._strip_json_code_block(negotiation_results_resp))
                for k, v in parsed.items():
                    doc.add_paragraph(f"{k}: {v}")
            except Exception:
                doc.add_paragraph(f"Raw: {negotiation_results_resp}")
        negotiation_results = self._handle_json_parse(negotiation_results_resp)
        if self._is_error(negotiation_results):
            print("API call failed for cross-cluster negotiation:", negotiation_results["__error__"])
            negotiation_results = {}
        if log_callback:
            log_callback("negotiation_results_parsed", negotiation_results)
    
        final_decision = self._conduct_weighted_vote_model2(problem, negotiation_results)
        if log_callback:
            log_callback("final_decision", asdict(final_decision))
    
        print(f"[DEBUG] considerations: {[asdict(c) for c in considerations]}")
    
        return final_decision
    

    def _execute_model3(self, problem: str, options, filtered_pairs, doc, considerations, log_callback=None) -> DecisionResult:
        if log_callback:
            log_callback("start", {"model": "model3", "problem": problem})
        print("Executing Model 3: Maslow-TA Decision Matrix")
        matrix_cells = []
        for ego_state in ["Parent", "Adult", "Child"]:
            for maslow_level, desc in [
                ("Physiological", "Basic survival needs: money, food, shelter, rest"),
                ("Safety", "Security, stability, health, protection from risk"),
                ("Love/Belonging", "Family, friends, community, social connection"),
                ("Esteem", "Status, recognition, mastery, achievement, respect"),
                ("Self-Actualization", "Purpose, growth, creativity, fulfilling potential")
            ]:
                prompt = self._matrix_cell_prompt(problem, ego_state, maslow_level, desc, options, filtered_pairs)
                response = self._call_openai_api(prompt)
    
                if log_callback:
                    log_callback("matrix_cell", {"ego_state": ego_state, "maslow_level": maslow_level, "raw_response": response})
                if doc:
                    doc.add_heading(f"{ego_state} x {maslow_level} Cell Analysis", level=3)
                    try:
                        parsed = json.loads(self._strip_json_code_block(response))
                        if isinstance(parsed, list):
                            for item in parsed:
                                doc.add_paragraph(f"{item.get('type', '')}: {item.get('text', '')} (Option: {item.get('option', '')}, Score: {item.get('score', '')})")
                        else:
                            for k, v in parsed.items():
                                doc.add_paragraph(f"{k}: {v}")
                    except Exception:
                        doc.add_paragraph(f"Raw: {response}")
    
                # Parse and extract considerations (expected as a list)
                data = self._handle_json_parse(response)
                if not self._is_error(data) and isinstance(data, list):
                    for cons in data:
                        c = Consideration(
                            text=cons.get("text", ""),
                            source_model="model3",
                            source_context=f"{ego_state}-{maslow_level}",
                            type=cons.get("type", ""),
                            option=cons.get("option"),
                            score=cons.get("score"),
                        )
                        considerations.append(c)
                        if log_callback:
                            log_callback("consideration", asdict(c))
                # You may wish to handle/report on errors or non-list responses here as appropriate.
    
                # MatrixCell for utility calculations, use the highest scored 'positive' or lowest 'negative' for each cell, or aggregate as desired
                # We'll use the "strongest" score for each cell for backward compatibility with utility calculation
                cell_score = 0.0
                cell_reasoning = ""
                if isinstance(data, list) and data:
                    # Prefer the first positive, otherwise the strongest
                    positive_scores = [float(item.get("score", 0)) for item in data if item.get("type") == "positive"]
                    if positive_scores:
                        cell_score = max(positive_scores)
                    else:
                        # fallback: get the most positive or least negative
                        all_scores = [float(item.get("score", 0)) for item in data]
                        if all_scores:
                            cell_score = max(all_scores, key=abs)
                    # Use the reasoning/text from the first consideration
                    cell_reasoning = data[0].get("text", "")
                elif isinstance(data, dict):
                    cell_score = float(data.get("score", 0))
                    cell_reasoning = data.get("text", "")
                else:
                    cell_score = 0.0
                    cell_reasoning = "API call failed or returned unexpected format."
                matrix_cells.append(
                    MatrixCell(
                        ego_state=ego_state,
                        maslow_level=maslow_level,
                        score=cell_score,
                        reasoning=cell_reasoning
                    )
                )
    
        if log_callback:
            log_callback("matrix_cells_collected", [asdict(cell) for cell in matrix_cells])
        tier_results, mitigations = self._check_tiers_model3(problem, matrix_cells)
        if log_callback:
            log_callback("tier_results", {"tier_results": tier_results, "mitigations": mitigations})
        if mitigations:
            for m in mitigations:
                print("Mitigation needed:", m)
                if doc:
                    doc.add_paragraph(f"Mitigation needed: {m}", style="Intense Quote")
        final_decision = self._calculate_utility_model3(problem, matrix_cells, tier_results)
        if log_callback:
            log_callback("final_decision", asdict(final_decision))
    
        print(f"[DEBUG] considerations: {[asdict(c) for c in considerations]}")
    
        return final_decision
    
    def execute_workflow(self, model: DecisionModel, problem: str, options, filtered_pairs, doc, considerations=None, log_callback=None) -> DecisionResult:
        if model == DecisionModel.DEMOCRATIC_COUNCIL:
            return self._execute_model1(problem, options, filtered_pairs, doc, considerations, log_callback)
        elif model == DecisionModel.SECOND_ORDER_NEGOTIATIONS:
            return self._execute_model2(problem, options, filtered_pairs, doc, considerations, log_callback)
        elif model == DecisionModel.MASLOW_TA_MATRIX:
            return self._execute_model3(problem, options, filtered_pairs, doc, considerations, log_callback)
        else:
            raise ValueError(f"Unknown model: {model}")

    # ==== Prompt construction helpers with numeric IDs ====

    def _ego_state_prompt(self, problem, ego_state, options, filtered_pairs):
        '''
        Constructs the prompt for Model 1 (Ego State Council) for a single ego state.
        Includes context: problem statement, pre-filtered option+outcome pairs.
        '''
        option_str = "\n".join([f"- {get_option_text(opt)}" for opt in options])
        context = self.format_context_for_prompt(problem, filtered_pairs)
        return f'''
{context}
You are analyzing this decision problem from the perspective of the {ego_state} ego state in Transactional Analysis.

Options under consideration:
{option_str}

For each option, identify up to three types of consideration:
1. **Positive Reason**: A reason to choose the option, based on a positive or desirable outcome if it is chosen.
2. **Negative Reason**: A reason to avoid the option, based on a negative or undesirable outcome if it is chosen.
3. **Avoidance/Preventative Reason**: A reason to choose the option as a safeguard, specifically because NOT choosing it would result in a negative or undesirable outcome.

**For each consideration, include:**
- "type": one of "positive", "negative", or "avoidance"
- "text": the explanation
- "option": the exact option text it refers to (copy from above)
- "score": from -10 (strongly against) to +10 (strongly for)

**Respond in valid JSON, as a flat list:**  
[
  {{"type": "positive", "text": "...", "option": "Option Text Here", "score": 8}},
  {{"type": "negative", "text": "...", "option": "Option Text Here", "score": -6}},
  {{"type": "avoidance", "text": "...", "option": "Option Text Here", "score": 7}}
]

Only include real, relevant considerations. Omit any type if there are no applicable reasons for that option.
'''



    def _council_synthesis_prompt(self, problem, ego_responses):
        return f'''
Synthesize the perspectives from all three ego states into a coherent council discussion.

Problem: {problem}

Ego State Responses:
{json.dumps([asdict(response) for response in ego_responses], indent=2)}

Identify shared themes, points of agreement, tension points, and propose mitigation plans.

Provide your synthesis in JSON format:
{{
    "shared_themes": ["list of themes all ego states agree on"],
    "tension_points": ["list of areas where ego states disagree"],
    "mitigation_plans": ["list of specific actions to address tensions"],
    "consensus_level": "score from 1-5 indicating how much agreement exists",
    "synthesis_reasoning": "detailed explanation of the council synthesis"
}}
'''

    def _sub_ego_state_desc(self, sub_state):
        sub_state_descriptions = {
            "Parent-in-Parent": "Moral authority and tradition, core values",
            "Adult-in-Parent": "Experienced rationality filtered through values",
            "Child-in-Parent": "Rules of conformity, following proper procedures",
            "Parent-in-Adult": "Structured problem-solving with ethical considerations",
            "Adult-in-Adult": "Pure data-driven analysis and logical reasoning",
            "Child-in-Adult": "Inventive curiosity within rational thinking",
            "Parent-in-Child": "Internalized rules about play and appropriate behavior",
            "Adult-in-Child": "Logical yet playful analysis, learning excitement",
            "Child-in-Child": "Pure spontaneity, emotion, and authentic desires"
        }
        return sub_state_descriptions.get(sub_state, "Unknown")

    def _sub_ego_state_prompt(self, problem, sub_state, options, filtered_pairs):
        '''
        Constructs the prompt for Model 2 (Sub-Ego States), one sub-ego at a time.
        Includes context: problem statement, pre-filtered option+outcome pairs.
        '''
        option_str = "\n".join([f"- {get_option_text(opt)}" for opt in options]) if options else ""
        context = self.format_context_for_prompt(problem, filtered_pairs)
        return f'''
{context}
You are analyzing this decision problem from the perspective of the {sub_state} sub-ego state in Second-Order Transactional Analysis.

Options under consideration:
{option_str}

{sub_state}: {self._sub_ego_state_desc(sub_state)}

For each option, identify up to three types of consideration:
1. **Positive Reason**: A reason to choose the option, based on a positive or desirable outcome if it is chosen.
2. **Negative Reason**: A reason to avoid the option, based on a negative or undesirable outcome if it is chosen.
3. **Avoidance/Preventative Reason**: A reason to choose the option as a safeguard, specifically because NOT choosing it would result in a negative or undesirable outcome.

**For each consideration, include:**
- "type": one of "positive", "negative", or "avoidance"
- "text": the explanation
- "option": the exact option text it refers to (copy from above)
- "score": from -10 (strongly against) to +10 (strongly for)

**Respond in valid JSON, as a flat list:**  
[
  {{"type": "positive", "text": "...", "option": "Option Text Here", "score": 8}},
  {{"type": "negative", "text": "...", "option": "Option Text Here", "score": -6}},
  {{"type": "avoidance", "text": "...", "option": "Option Text Here", "score": 7}}
]
'''


    def _cluster_dialogues_prompt(self, problem, sub_ego_responses):
        return f'''
Conduct internal dialogues within clusters to reach agreement.

Problem: {problem}

Cluster Members:
{json.dumps([asdict(member) for member in sub_ego_responses], indent=2)}

Find common ground and negotiate stance changes within clusters.

Provide results in JSON format (group by clusters as needed):
{{
    "cluster_name": "name",
    "main_negotiation_points": ["key issues discussed"],
    "agreed_actions": ["specific agreements reached"],
    "stance_changes": [
        {{"sub_state": "name", "old_stance": 0,
            "new_stance": 1, "reason": "explanation"}}
    ],
    "cluster_consensus": "score from 1-5"
}}
'''

    def _cross_cluster_negotiation_prompt(self, problem, cluster_results):
        return f'''
Conduct cross-cluster negotiations to resolve remaining disagreements.

Problem: {problem}

Cluster Results:
{json.dumps(cluster_results, indent=2)}

Identify remaining concerns and create horse-trading proposals.

Provide results in JSON format:
{{
    "pending_concerns": ["issues still needing resolution"],
    "trade_proposals": [
        {{"concern": "issue", "blocking_voice": "sub-state",
            "proposal": "solution", "result": "outcome"}}
    ],
    "final_stance_changes": [
        {{"sub_state": "name", "old_stance": 0, "new_stance": 1}}
    ],
    "negotiation_success": "score from 1-5"
}}
'''

    def _matrix_cell_prompt(self, problem, ego_state, maslow_level, maslow_desc, options, filtered_pairs):
        '''
        Constructs the prompt for Model 3 (Maslow-TA Matrix) for one cell (ego state x maslow level).
        Includes context: problem statement, pre-filtered option+outcome pairs.
        '''
        option_str = "\n".join([f"- {get_option_text(opt)}" for opt in options]) if options else ""
        context = self.format_context_for_prompt(problem, filtered_pairs)
        return f'''
{context}
Evaluate how the decision impacts {maslow_level} needs from the {ego_state} ego state perspective.

Options under consideration:
{option_str}

{ego_state} Perspective:
- Parent: Values, morals, social expectations, duty
- Adult: Facts, logic, practical consequences
- Child: Emotions, desires, authenticity, spontaneity

{maslow_level}: {maslow_desc}

For each option, identify up to three types of consideration:
1. **Positive Reason**: A reason to choose the option, based on a positive or desirable outcome if it is chosen.
2. **Negative Reason**: A reason to avoid the option, based on a negative or undesirable outcome if it is chosen.
3. **Avoidance/Preventative Reason**: A reason to choose the option as a safeguard, specifically because NOT choosing it would result in a negative or undesirable outcome.

For each consideration, include:
- "type": one of "positive", "negative", or "avoidance"
- "text": the detailed explanation
- "option": the exact option text it refers to (copy from above)
- "score": from -10 (strongly against) to +10 (strongly for)

Respond in valid JSON, as a flat list:
[
  {{"type": "positive", "text": "...", "option": "Option Text Here", "score": 7}},
  {{"type": "avoidance", "text": "...", "option": "Option Text Here", "score": 5}},
  ...
]
'''

    def _summarize_model_output_prompt(self, problem: str, model_output: Dict[str, Any]) -> str:
        '''
        Generates a prompt for the AI to summarize the model's output in Markdown.
        '''
        return f'''
The following is the full output from one of the Transactional Analysis decision models.
Your task is to provide a concise and insightful summary of the model's findings,
recommendation, and reasoning in Markdown format.

Original Decision Problem: {problem}

Model Output:
```json
{json.dumps(model_output, indent=2)}
Please summarize this information, highlighting the key recommendation, confidence,
and any important conditions or detailed analysis. Format your response using Markdown.
'''

    @staticmethod
    def format_context_for_prompt(problem, filtered_pairs):
        '''
        Builds a plain English block describing the original question, and all currently valid
        (option, likely outcome) pairs, for injection at the top of each model's prompt.
        '''
        pairs = "\n".join(
            [f'- Option: "{item["option_text"]}"\n  Likely outcome: "{item["likely_outcome"]}"'
             for item in filtered_pairs]
        )
        return f'''Decision problem: {problem}

Consider ONLY the following pre-validated options and their likely outcomes for your analysis:
{pairs}
'''
    

####################################################################################################
# Consideration Storage & Indexing
####################################################################################################
from collections import defaultdict
from typing import List, Dict, Optional
import difflib

class ConsiderationProcessor:
    '''
    In-memory indexed storage for all considerations produced by model runs.
    Supports grouping, de-duplication, lookup by option/model/context, and full traceability.
    Used to prepare output tables and cross-reference model results.
    '''
    def __init__(self, logger=print):
        self._table = []  # All considerations as list of dataclass or dict
        self._by_option_text = {}  # str: list[Consideration]
        self._by_option_id = {}    # str: list[Consideration]
        self._by_model = {}        # model: list[Consideration]
        self._by_context = {}      # context: list[Consideration]
        self._general = []         # Not tied to any option
        self.logger = logger

        # Canonical option lists/mapping (populated via set_options)
        self._option_texts = []      # List of canonical option texts
        self._option_id_by_text = {} # Dict: canonical text → ID

    def set_options(self, options: list):
        self._option_texts = [get_option_text(opt) for opt in options]
        self._option_id_by_text = {get_option_text(opt): opt['id'] for opt in options}
        self.logger(f"[ConsiderationDB][set_options] Canonical options: {self._option_texts}")

    def resolve_option(self, raw_option_text):
        if not raw_option_text:
            return None
        if raw_option_text in self._option_texts:
            return raw_option_text
        matches = difflib.get_close_matches(raw_option_text, self._option_texts, n=1, cutoff=0.8)
        if matches:
            self.logger(f"[ConsiderationDB][resolve_option] Fuzzy-matched '{raw_option_text}' to '{matches[0]}'")
            return matches[0]
        self.logger(f"[ConsiderationDB][resolve_option] No match for '{raw_option_text}'")
        return None

    def add(self, cons):
        print("[DEBUG][add] Adding consideration:", cons)
        self._table.append(cons)
        print(f"[DEBUG][add] _table now has {len(self._table)} items.")
        option_text = getattr(cons, 'option', None)
        option_id = getattr(cons, 'option_id', None)
        canonical_option = None
        print(f"[DEBUG][add] option_text: {option_text}, option_id: {option_id}")
        if option_text:
            canonical_option = self.resolve_option(option_text)
            print(f"[DEBUG][add] canonical_option resolved: {canonical_option}")
            if canonical_option:
                cons.option = canonical_option
                self._by_option_text.setdefault(canonical_option, []).append(cons)
                print(f"[DEBUG][add] Added to _by_option_text under {canonical_option}.")
            else:
                self._general.append(cons)
                print(f"[DEBUG][add] Option text not resolved, added to _general.")
        else:
            self._general.append(cons)
            print(f"[DEBUG][add] No option_text, added to _general.")
        if option_id:
            self._by_option_id.setdefault(option_id, []).append(cons)
            print(f"[DEBUG][add] Added to _by_option_id under {option_id}.")
        model = getattr(cons, 'source_model', None)
        if model:
            self._by_model.setdefault(model, []).append(cons)
            print(f"[DEBUG][add] Added to _by_model under {model}.")
        context = getattr(cons, 'source_context', None)
        if context:
            self._by_context.setdefault(context, []).append(cons)
            print(f"[DEBUG][add] Added to _by_context under {context}.")
        self.logger(
            f"[ConsiderationDB][add] Added: model={model}, option_text={option_text}, "
            f"canonical_option={canonical_option}, option_id={option_id}, "
            f"context={context}, orientation={getattr(cons,'orientation',None)}, "
            f"text={getattr(cons,'text','')[:90]}"
        )
        print(f"[DEBUG][add] Completed processing for: {getattr(cons, 'text', '')[:40]}...")

    def add_many(self, conslist):
        print(f"[ConsiderationDB][add_many] Called with {len(conslist)} considerations")
        if not conslist:
            print("[ConsiderationDB][add_many] Warning: Empty list provided.")
        for i, c in enumerate(conslist):
            print(f"[ConsiderationDB][add_many] Adding consideration {i+1}/{len(conslist)}: {getattr(c, 'text', repr(c))[:80]}")
            self.add(c)
        print(f"[ConsiderationDB][add_many] Finished adding all considerations.")

    def all(self):
        return list(self._table)

    def by_option_text(self, option_text):
        cons = self._by_option_text.get(option_text, [])
        self.logger(f"[ConsiderationDB][by_option_text] Queried '{option_text}' -> {len(cons)} found.")
        return cons

    def by_option_id(self, option_id):
        cons = self._by_option_id.get(option_id, [])
        self.logger(f"[ConsiderationDB][by_option_id] Queried '{option_id}' -> {len(cons)} found.")
        return cons

    def by_model(self, model):
        cons = self._by_model.get(model, [])
        self.logger(f"[ConsiderationDB][by_model] Queried '{model}' -> {len(cons)} found.")
        return cons

    def by_context(self, context):
        cons = self._by_context.get(context, [])
        self.logger(f"[ConsiderationDB][by_context] Queried '{context}' -> {len(cons)} found.")
        return cons

    def general(self):
        self.logger(f"[ConsiderationDB][general] Queried -> {len(self._general)} found.")
        return list(self._general)

    def as_table_data(self, conslist):
        return [
            [getattr(c, 'text', ''), getattr(c, 'orientation', ''), str(getattr(c, 'score', ''))]
            for c in conslist
        ]

    def clear(self):
        self._table.clear()
        self._by_option_text.clear()
        self._by_option_id.clear()
        self._by_model.clear()
        self._by_context.clear()
        self._general.clear()
        self.logger("[ConsiderationDB][clear] All data cleared.")

    def dump_summary(self):
        self.logger(f"[ConsiderationDB][dump_summary] {len(self._table)} total considerations.")
        for key, dic in [
            ('option_text', self._by_option_text),
            ('model', self._by_model),
            ('context', self._by_context)
        ]:
            self.logger(f"  - {key}: {list(dic.keys())}")

    

# ========== MAIN DRIVER ==========
# --- Model synopses ---
model_synopses = {
    "model1": (
    "Model 1: Democratic Ego State Council\n"
    "This model simulates a decision as an internal council between the three core ego states of Transactional Analysis: Parent, Adult, and Child. "
    "Each ego state independently considers the problem, then the responses are synthesized into a council discussion. "
    "The final recommendation is based on consensus or majority among the three voices."
),
    "model2": (
    "Model 2: Second-Order Ego State Negotiations\n"
    "This model considers the nuanced perspectives of nine sub-ego states (Parent-in-Parent, Adult-in-Parent, etc.). "
    "It mimics a more complex negotiation process with internal 'clusters' and cross-cluster horse-trading, before arriving at a weighted consensus."
),
    "model3": (
    "Model 3: Maslow-TA Decision Matrix\n"
    "This model evaluates the decision's impact on each ego state, tiered by the levels of Maslow’s Hierarchy of Needs. "
    "The matrix approach highlights where core needs are at risk, and calculates an overall utility score to drive the recommendation."
)
}

def add_spoken_synopsis_to_doc(doc, results):
    '''
    Adds a narrative explanation of each model’s confidence score and result to the Word doc.
    Intended for more accessible/layperson understanding in the report.
    '''
    # Add the spoken-out explanation of each model's confidence score to the docx report.
    doc.add_page_break()
    doc.add_heading("Final Model Comparison & Synopsis", level=1)
    for model_name, result in results.items():
        doc.add_heading(result.model_used, level=2)
        doc.add_paragraph(f"Recommendation: {result.recommendation}")

        if model_name == 'model1':
            avg = result.confidence_score  # -10 to +10
            doc.add_paragraph(
                f"Raw Confidence Score: {avg:.2f} (on a scale of -10 to +10)"
            )
            # Normalized confidence for comparison (0 to 1)
            norm = (avg + 10) / 20.0
            doc.add_paragraph(
                f"Normalized Confidence: {norm:.2f} (scale 0 to 1)"
            )
            # Conditional explanation
            if avg >= 7.5:
                detail = "This indicates a very strong positive consensus among the three ego states. The decision is almost unanimously supported."
            elif avg >= 5.0:
                detail = "This means a strong consensus in favor. Most ego states are clearly for this decision."
            elif avg >= 2.5:
                detail = "This suggests a moderately positive consensus. The group generally supports the decision, but there may be some reservations."
            elif avg > 0:
                detail = "This is a neutral to slightly positive score. The council isn't opposed, but also isn't strongly convinced."
            elif avg == 0:
                detail = "This is a perfectly balanced score. The council is evenly split."
            elif avg > -2.5:
                detail = "This means a mildly negative consensus. The decision is not recommended, but not strongly rejected."
            elif avg > -5.0:
                detail = "This indicates a moderately negative consensus. Most ego states are hesitant or opposed."
            else:
                detail = "This is a strong rejection. Nearly all ego states are against this decision."
            doc.add_paragraph(
                "The Democratic Council model calculates confidence as the average of all ego state votes, with -10 being strongly against, 0 perfectly neutral, and +10 strongly for. " + detail)

        elif model_name == 'model2':
            conf = result.confidence_score  # -1 to +1
            doc.add_paragraph(f"Weighted Confidence: {conf:.2f} (scale -1 to +1)")
            if conf >= 0.8:
                detail = "There is robust agreement among the sub-ego states—clear consensus."
            elif conf >= 0.5:
                detail = "There is broad support, but some disagreement exists."
            elif conf >= 0.1:
                detail = "The outcome is mixed; support is present, but notable dissent remains."
            elif conf == 0:
                detail = "The group is evenly split; acceptance and rejection are balanced."
            elif conf > -0.1:
                detail = "The group is mixed, but slightly tilted negative; rejection is weakly justified."
            elif conf > -0.5:
                detail = "The group is mostly negative; most sub-ego states oppose the decision."
            else:
                detail = "There is no support; the group is unified in rejection."
            doc.add_paragraph(
                "The Second-Order Negotiations model computes a weighted average of sub-ego states' stances, with -1 meaning unanimous rejection, 0 balanced, and +1 unanimous support. " + detail)

        elif model_name == 'model3':
            conf = result.confidence_score  # 0 to 1, with 0.5 as neutral
            doc.add_paragraph(
                f"Utility Score: {conf:.2f} (scale 0 to 1, with 0.5 as neutral, higher is better)"
            )
            if conf >= 0.85:
                detail = "This decision strongly fulfills all levels of Maslow's needs. It is highly beneficial."
            elif conf >= 0.7:
                detail = "The decision fulfills most needs well; only minor issues exist at certain levels."
            elif conf >= 0.55:
                detail = "The decision is adequate; most needs are met, but there are notable areas for improvement."
            elif conf > 0.5:
                detail = "The decision is slightly positive; most core needs are met, though with reservations."
            elif conf == 0.5:
                detail = "The decision is net neutral; benefits and risks are balanced."
            elif conf > 0.4:
                detail = "The decision is risky or marginal; key needs may be left unmet."
            else:
                detail = "The decision fails to meet essential needs at several levels; it is not recommended."
            doc.add_paragraph(
                "The Maslow-TA Matrix model reflects how well the choice satisfies all layers of psychological and practical needs. 0.5 means perfectly neutral (no overall gain or loss), 1 is maximum fulfillment, 0 is maximum risk/negativity. " + detail)

        doc.add_paragraph(f"Conditions/Notes: {result.conditions}")
        doc.add_paragraph(f"Summary Reasoning: {result.reasoning}")


####################################################################################################
# Main Report Generation Driver
####################################################################################################
def main():
    '''
    Main entry point for the Decisionator tool.
    '''
    if OPENAI_API_KEY == "your-openai-api-key-here":
        print("ERROR: Please set your OpenAI API key in the OPENAI_API_KEY variable")
        return

    if DEBUG_MODE:
        test_line = "* **Financial Urgency:** The family needs significant income soon to pay for his daughter's school fees (due in 2 weeks: ~$300 USD equivalent) and general living expenses."
        debug_process_specific_md_line(test_line)
    

    # Prompt user for problem
    print("Please enter the decision problem you want to analyze.")
    print("To finish, press Enter twice on consecutive empty lines (i.e., press Enter, then press Enter again):")
    problem_lines = []
    empty_line_count = 0
    while True:
        line = input()
        if not line:
            empty_line_count += 1
            if empty_line_count >= 2:
                break
        else:
            empty_line_count = 0
            problem_lines.append(line)
    problem = "\n".join(problem_lines).strip()
    while not problem:
        print("The question cannot be blank. Please enter a valid question.")
        print("Please enter the decision problem you want to analyze.")
        print("To finish, press Enter twice on consecutive empty lines (i.e., press Enter, then press Enter again):")
        problem_lines = []
        empty_line_count = 0
        while True:
            line = input()
            if not line:
                empty_line_count += 1
                if empty_line_count >= 2:
                    break
            else:
                empty_line_count = 0
                problem_lines.append(line)
        problem = "\n".join(problem_lines).strip()

    docasm = DocAssembler()
    controller = WorkflowController()
    processor = ConsiderationProcessor()

    # ==== Problem and Options Section ====
    docasm.add_heading("Problem Statement", level=1)
    docasm.add_heading("Original Question", level=2)
    if is_markdown(problem):
        docasm.add_markdown(problem)
    else:
        docasm.add_paragraph(problem.strip())

    # ==== Extract options ====
    print("\nDetecting decision options in the problem statement...")
    options = controller.extract_options(problem)
    if not options:
        print("No options were extracted from the problem. Please rephrase your problem or check your input.")
        sys.exit(1)
    print("Options detected:")
    for opt in options:
        print(f"{opt['id']}: {get_option_text(opt)}")

    # === After extracting options, determine likely outcomes for each option ===
    likely_outcomes = []
    for opt in options:
        prompt = f'''
You are an expert in scenario analysis. Your task is to predict the **most likely realistic outcome** for the given decision option.

**Instructions:**
- Carefully read the full problem statement.
- Take into account all factors, constraints, and context provided in the problem.
- Use your own general knowledge and reasoning to augment or clarify the scenario as needed.
- Think through any consequences or side effects, considering both direct and indirect effects.
- Your outcome should represent what *actually* would most probably happen, not an ideal or worst-case, but the typical or most expected result.
- Be specific and practical.

**Option:** "{get_option_text(opt)}"
**Full Problem Context:** "{problem}"

Return as JSON with fields:
{{
  "option_id": "{opt['id']}",
  "option_text": "{get_option_text(opt)}",
  "likely_outcome": "<provide the most probable, concrete outcome here>"
}}
'''
        response = controller._call_openai_api(prompt)
        data = controller._handle_json_parse(response)
        likely_outcomes.append(data)

    print("\nLikely outcomes for each option:")
    for item in likely_outcomes:
        print(f"- {item.get('option_text', '[MISSING]')}: {item.get('likely_outcome', '[NO OUTCOME]')}")

    # === Appraise (option, outcome) pairs for deleterious results ===
    filtered_pairs = []
    rejection_reasons = []
    for item in likely_outcomes:
        prompt = f'''
Evaluate the following decision context. If choosing this option and the likely outcome as described would result in death or injury (to self or others), a logical impossibility, or breaking the law, say so.
Context:
Option: "{item['option_text']}"
Likely outcome: "{item['likely_outcome']}"
Return JSON:
{{"is_deleterious": true/false, "reason": "<explanation>"}}
'''
        resp = controller._call_openai_api(prompt)
        check = controller._handle_json_parse(resp)
        if check.get("is_deleterious"):
            rejection_reasons.append({
                "option_text": item['option_text'],
                "likely_outcome": item['likely_outcome'],
                "reason": check.get("reason", "")
            })
        else:
            filtered_pairs.append(item)

    # If all are rejected, print explanations and abort:
    if not filtered_pairs:
        print("\nAll possible courses of action and their likely outcomes were rejected due to deleterious consequences:")
        for rej in rejection_reasons:
            print(f"- Option: {rej['option_text']}\n  Outcome: {rej['likely_outcome']}\n  Reason: {rej['reason']}\n")
        print("No further decision processing can continue for this problem.")
        sys.exit(0)

    print("\nSurviving (option, likely outcome) pairs:")
    for item in filtered_pairs:
        print(f"- {item.get('option_text', '[MISSING]')}: {item.get('likely_outcome', '[NO OUTCOME]')}")

    docasm.add_heading("Options Considered", level=2)
    for opt in options:
        docasm.add_paragraph(f"{opt['id']}: {get_option_text(opt)}")

    processor.set_options(options)

    # ==== MODEL SELECTION MENU ====
    print("\nTA Decision-Making Models Menu")
    print("=" * 50)
    print("Which decision model would you like to run?")
    print("  1: Democratic Ego State Council")
    print("     - Simulates a council between the Parent, Adult, and Child ego states, synthesizing a consensus.\n")
    print("  2: Second-Order Ego State Negotiations")
    print("     - Considers nine sub-ego states and mimics a complex negotiation with weighted consensus.\n")
    print("  3: Maslow-TA Decision Matrix")
    print("     - Evaluates how each option meets the levels of Maslow’s Hierarchy of Needs, using a utility score.\n")
    print("  4: All models (recommended for full analysis)")

    model_choice = None
    valid_choices = {"1", "2", "3", "4"}
    while True:
        inp = input("Enter the number (1, 2, 3, or 4) of the model to run [4=all]: ").strip()
        if inp in valid_choices:
            if inp == "4":
                model_choice = "all"
            else:
                model_choice = int(inp)
            break
        print("Invalid input. Please enter 1, 2, 3, or 4.")

    if model_choice == "all":
        models_to_run = [
            DecisionModel.DEMOCRATIC_COUNCIL,
            DecisionModel.SECOND_ORDER_NEGOTIATIONS,
            DecisionModel.MASLOW_TA_MATRIX
        ]
    elif model_choice == 1:
        models_to_run = [DecisionModel.DEMOCRATIC_COUNCIL]
    elif model_choice == 2:
        models_to_run = [DecisionModel.SECOND_ORDER_NEGOTIATIONS]
    elif model_choice == 3:
        models_to_run = [DecisionModel.MASLOW_TA_MATRIX]
    else:
        print("No valid model selected, exiting.")
        sys.exit(1)

    print("\nRunning the following models:")
    for m in models_to_run:
        print(" -", m.name.replace("_", " ").title())

    results = {}
    everyday_summaries = {}
    steps_by_model = {}

    # ==== Considerations by Option (Chapter 3) ====
    docasm.add_page_break()
    docasm.add_heading("Decision Considerations (by Option)", level=1)
    docasm.add_markdown(
        "The considerations listed below are grouped by decision option and categorized as follows:\n"
        "- **Negative:** Arguments against choosing this option, based on possible negative outcomes if it is chosen. These are typically shown with negative scores.\n"
        "- **Avoidance:** Arguments in favor of this option, specifically because *not* choosing it would lead to negative consequences. These are preventative or risk-avoiding reasons and are scored positively.\n"
        "- **Positive:** Arguments directly in favor of this option, based on desirable outcomes if it is chosen. These also have positive scores.\n\n"
        "Where several considerations are highly similar, they are grouped together and summarized with an introductory sentence and bullet points."
    )

    # --- Run models and collect considerations ---
    for i, model in enumerate(models_to_run):
        model_label = model.value
        print(f"\n{'='*10} Running Model {i+1}: {model.name.replace('_', ' ').title()} {'='*10}")
        model_steps = []
        def log_step(step_name, data):
            model_steps.append({'step': step_name, 'data': data})
        temp_considerations = []
        result = controller.execute_workflow(
            model,
            problem,
            options,
            filtered_pairs,
            doc=None,
            considerations=temp_considerations,
            log_callback=log_step
        )
        processor.add_many(temp_considerations)
        results[model_label] = result
        steps_by_model[model_label] = model_steps

        everyday_prompt = controller._everyday_language_summary_prompt(problem, asdict(result))
        everyday_summary = controller._call_openai_api(everyday_prompt)
        everyday_summaries[model_label] = everyday_summary

    # -- Group considerations by option --
    grouped = {get_option_text(opt): processor.by_option_text(get_option_text(opt)) for opt in options}
    general_cons = processor.general()
    all_cons = processor.all()

    # ---- Compute normalization for all considerations ----
    neg_scores = [float(c.score) for c in all_cons if c.type == "negative" and c.score is not None and float(c.score) < 0]
    pos_scores = [float(c.score) for c in all_cons if c.type in ("positive", "avoidance") and c.score is not None and float(c.score) > 0]
    min_neg = min(neg_scores) if neg_scores else -1
    max_pos = max(pos_scores) if pos_scores else 1

    # Build a mapping from option_text to rejection reason for easy lookup
    rejection_by_option = {rej['option_text']: rej for rej in rejection_reasons}

    for opt in options:
        opt_text = get_option_text(opt)
        conslist = grouped.get(opt_text, [])
        conslist = dedupe_considerations(conslist, similarity_cutoff=0.95)

        docasm.add_heading(f"{opt['id']}: {get_option_text(opt)}", level=2)

        # If this option was rejected, print its rejection reason and skip table
        if opt_text in rejection_by_option:
            rej = rejection_by_option[opt_text]
            docasm.add_markdown(
                f"This option was **rejected and excluded from further analysis** because:\n\n"
                f"{rej['reason']}\n\n"
                f"**Likely outcome if chosen:** {rej['likely_outcome']}"
            )
            continue

        if not conslist:
            docasm.add_paragraph("No considerations for this option.")
            continue

        # --- NEW AI-BASED GROUPING AND MERGING ---
        merged_groups = merge_considerations_conceptually(conslist, controller, context_prompt=problem)
        table = docasm.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Consideration'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Score'
        style_table(table)

        for group in merged_groups:
            row_cells = table.add_row().cells
            docasm.add_markdown_to_cell(row_cells[0], group['merged_text'])
            row_cells[1].text = str(group['type'])
            row_cells[2].text = f"{group['score']:.2f}" if group['score'] is not None else ''
            try:
                s = float(group['score'])
            except Exception:
                s = 0
            if s < 0 or s > 0:
                hexcolor = get_heatmap_color(s, min_neg, max_pos, group['type'])
                if hexcolor:
                    for cell in row_cells:
                        cell._tc.get_or_add_tcPr().append(
                            parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hexcolor))
                        )

    # --- General considerations as usual ---
    docasm.add_heading("General Considerations (Not Tied to a Single Option)", level=2)
    general_cons_short = dedupe_considerations(general_cons, similarity_cutoff=0.95)
    if general_cons_short:
        merged_groups = merge_considerations_conceptually(general_cons_short, controller, context_prompt=problem)
        table = docasm.add_table(rows=1, cols=3)
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'Consideration'
        hdr_cells[1].text = 'Type'
        hdr_cells[2].text = 'Score'
        style_table(table)

        for group in merged_groups:
            row_cells = table.add_row().cells
            docasm.add_markdown_to_cell(row_cells[0], group['merged_text'])
            row_cells[1].text = str(group['type'])
            row_cells[2].text = f"{group['score']:.2f}" if group['score'] is not None else ''
            try:
                s = float(group['score'])
            except Exception:
                s = 0
            if s < 0 or s > 0:
                hexcolor = get_heatmap_color(s, min_neg, max_pos, group['type'])
                if hexcolor:
                    for cell in row_cells:
                        cell._tc.get_or_add_tcPr().append(
                            parse_xml(r'<w:shd {} w:fill="{}"/>'.format(nsdecls('w'), hexcolor))
                        )
    else:
        docasm.add_paragraph("No general considerations were raised.")

    # ==== Chapter 2: Summaries and Conclusions ====
    docasm.add_page_break()
    docasm.add_heading("Summary Conclusions and Recommendations", level=1)
    docasm.add_paragraph(
        "This section collects each model’s everyday-language summary and provides the AI's overall recommendation."
    )

    for model_label in [m.value for m in models_to_run]:
        docasm.add_heading(f"{model_label.upper()} Model Summary", level=2)
        docasm.add_markdown(everyday_summaries[model_label])

    # Build the overall summary prompt dynamically (robust for any model subset)
    summary_lines = [
        "Here is a decision problem and the following model output summaries, each written in plain, everyday language.",
        "Your job is to give a clear, everyday-language overall conclusion and recommendation, comparing the model outputs provided below. Explain which model's advice makes the most sense in simple, practical terms.",
        "",
        f"Original Question:\n{problem}",
        ""
    ]
    for model_label in [m.value for m in models_to_run]:
        model_name_nice = {
            "model1": "Model 1: Democratic Ego State Council",
            "model2": "Model 2: Second-Order Ego State Negotiations",
            "model3": "Model 3: Maslow-TA Decision Matrix"
        }.get(model_label, model_label.upper())
        summary_lines.append(f"{model_name_nice} Summary:\n{everyday_summaries[model_label]}\n")
    summary_lines.append(
        "**Please give your conclusion and recommendation for a general audience, using clear narrative or bullet points in Markdown. Do not use JSON or code blocks. Only refer to the models and summaries given above.**"
    )
    overall_prompt = "\n".join(summary_lines)
    overall_summary = controller._call_openai_api(overall_prompt)
    docasm.add_heading("Overall Conclusion and Recommendation", level=2)
    docasm.add_markdown(overall_summary)

    ##Uncomment the following code block if detailed working for each model is required.
    '''
    # --------- Appendix with Detailed Model Output (NO summary conclusions here) ---------
    docasm.add_page_break()
    docasm.add_heading("Appendix", level=1)
    docasm.add_paragraph(
        "The findings presented in the main part of the document are summarised from the following detailed output from each decision model."
    )

    for i, model in enumerate(models_to_run):
        model_label = model.value
        docasm.add_heading(f"Model: {model_label.upper()}", level=2)
        if model_label in globals().get('model_synopses', {}):
            docasm.add_paragraph(model_synopses[model_label], style='Intense Quote')
        docasm.add_heading("Workthrough", level=3)
        for step in steps_by_model[model_label]:
            docasm.add_heading(f"Step: {step['step']}", level=4)
            if isinstance(step['data'], dict):
                for k, v in step['data'].items():
                    docasm.add_paragraph(f"{k}: {v}")
            else:
                docasm.add_paragraph(str(step['data']))
        docasm.add_heading("Model Result", level=3)
        res = results[model_label]
        docasm.add_paragraph(f"Recommendation: {res.recommendation}")
        docasm.add_paragraph(f"Confidence Score: {res.confidence_score:.2f}")
        docasm.add_paragraph(f"Conditions: {res.conditions}")
        docasm.add_paragraph(f"Reasoning: {res.reasoning}")
        summary_prompt = controller._summarize_model_output_prompt(problem, asdict(res))
        summary_markdown = controller._call_openai_api(summary_prompt)
        docasm.add_heading("AI Summary of Model Results", level=4)
        docasm.add_markdown(summary_markdown)
        if i < len(models_to_run) - 1:
            docasm.add_page_break()
    '''

    # --------- Save document ---------
    timestamp = datetime.datetime.now().strftime("%Y%m%d_%H%M%S")
    output_filename = f"TA_Decision_Report_[{timestamp}].docx"
    docasm.save(output_filename)
    print(f"\nReport saved as {output_filename}")
    return results

def explain_results_speech(results):
    '''
    Print a spoken-out explanation of each model's confidence score and recommendation,
    and explain what each particular score denotes.
    '''
    for model_name, result in results.items():
        print("=" * 60)
        print(f"Model: {result.model_used}")
        print(f"Recommendation: {result.recommendation}")

        if model_name == 'model1':
            avg = result.confidence_score  # -10 to +10
            norm = (avg + 10) / 20.0      # Normalized to 0..1
            print(f"Raw Confidence Score: {avg:.2f} (on a scale of -10 to +10)")
            print(f"Normalized Confidence: {norm:.2f} (scale 0 to 1)")
            # Detailed conditional explanation:
            if avg >= 7.5:
                detail = "This indicates a *very strong* positive consensus among the three ego states. The decision is almost unanimously supported."
            elif avg >= 5.0:
                detail = "This means a *strong* consensus in favor. Most ego states are clearly for this decision."
            elif avg >= 2.5:
                detail = "This suggests a *moderately positive* consensus. The group generally supports the decision, but there may be some reservations."
            elif avg > 0:
                detail = "This is a *neutral* to slightly positive score. The council isn't opposed, but also isn't strongly convinced."
            elif avg == 0:
                detail = "This is a *perfectly balanced* score. The council is evenly split."
            elif avg > -2.5:
                detail = "This means a *mildly negative* consensus. The decision is not recommended, but not strongly rejected."
            elif avg > -5.0:
                detail = "This indicates a *moderately negative* consensus. Most ego states are hesitant or opposed."
            else:
                detail = "This is a *strong rejection*. Nearly all ego states are against this decision."
            print(
                "The Democratic Council model calculates confidence as the average of all ego state votes, "
                "with -10 being strongly against, 0 perfectly neutral, and +10 strongly for. " + detail
            )

        elif model_name == 'model2':
            conf = result.confidence_score  # -1 to +1
            print(f"Weighted Confidence: {conf:.2f} (scale -1 to +1)")
            if conf >= 0.8:
                detail = "There is *robust agreement* among the sub-ego states—clear consensus."
            elif conf >= 0.5:
                detail = "There is *broad support*, but some disagreement exists."
            elif conf >= 0.1:
                detail = "The outcome is *mixed*; support is present, but notable dissent remains."
            elif conf == 0:
                detail = "The group is *evenly split*; acceptance and rejection are balanced."
            elif conf > -0.1:
                detail = "The group is *mixed*, but slightly tilted negative; rejection is weakly justified."
            elif conf > -0.5:
                detail = "The group is *mostly negative*; most sub-ego states oppose the decision."
            else:
                detail = "There is *no support*; the group is unified in rejection."
            print(
                "The Second-Order Negotiations model computes a weighted average of sub-ego states' stances, "
                "with -1 meaning unanimous rejection, 0 balanced, and +1 unanimous support. " + detail
            )

        elif model_name == 'model3':
            conf = result.confidence_score  # 0 to 1, with 0.5 as neutral
            print(f"Utility Score: {conf:.2f} (scale 0 to 1, with 0.5 as neutral, higher is better)")
            if conf >= 0.85:
                detail = "This decision *strongly* fulfills all levels of Maslow's needs. It is highly beneficial."
            elif conf >= 0.7:
                detail = "The decision fulfills most needs *well*; only minor issues exist at certain levels."
            elif conf >= 0.55:
                detail = "The decision is *adequate*; most needs are met, but there are notable areas for improvement."
            elif conf > 0.5:
                detail = "The decision is *slightly positive*; most core needs are met, though with reservations."
            elif conf == 0.5:
                detail = "The decision is *net neutral*; benefits and risks are balanced."
            elif conf > 0.4:
                detail = "The decision is *risky or marginal*; key needs may be left unmet."
            else:
                detail = "The decision fails to meet *essential needs* at several levels; it is not recommended."
            print(
                "The Maslow-TA Matrix model reflects how well the choice satisfies all layers of psychological and practical needs. "
                "0.5 means perfectly neutral (no overall gain or loss), 1 is maximum fulfillment, 0 is maximum risk/negativity. " + detail
            )
        print(f"Conditions/Notes: {result.conditions}")
        print(f"Summary Reasoning: {result.reasoning}")
        print()


####################################################################################################
# Script Entry Point
####################################################################################################
if __name__ == "__main__":
    main()
